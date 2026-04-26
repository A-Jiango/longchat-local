from __future__ import annotations

import csv
import hashlib
import io
import math
import os
import re
import shutil
import subprocess
import sys
import tempfile
import time
import zipfile
from collections import OrderedDict
from concurrent.futures import Future, ThreadPoolExecutor
from dataclasses import dataclass
from html.parser import HTMLParser
from pathlib import Path
from xml.etree import ElementTree

from PySide6.QtCore import QByteArray, QEvent, QObject, QPoint, QRect, QRectF, QSize, QThread, QTimer, Qt, Signal, Slot
from PySide6.QtGui import QCloseEvent, QColor, QFont, QIcon, QPainter, QPen, QPixmap, QTextDocument, QTextOption
try:
    from PySide6.QtSvg import QSvgRenderer
except Exception:
    QSvgRenderer = None
from PySide6.QtWidgets import (
    QAbstractSlider,
    QApplication,
    QFileDialog,
    QFrame,
    QGraphicsDropShadowEffect,
    QHBoxLayout,
    QLabel,
    QLayout,
    QLayoutItem,
    QMainWindow,
    QMessageBox,
    QPlainTextEdit,
    QPushButton,
    QScrollArea,
    QSizePolicy,
    QStackedLayout,
    QVBoxLayout,
    QWidget,
)


WINDOW_TITLE = "Local Chat"
HEADER_TITLE = "LLM"
BRAND_BLUE = "#1677FF"
NEW_CHAT_ICON = Path(__file__).with_name("assets") / "new_chat.svg"
INPUT_PANEL_WIDTH_RATIO = 0.75
COMPOSER_MAX_VISIBLE_LINES = 3
ASSISTANT_RENDER_INTERVAL_MS = 60
STREAMING_RENDER_INTERVAL_MS = 16
DISPLAY_STREAM_INTERVAL_MS = 16
DISPLAY_STREAM_MIN_UNITS = 1
DISPLAY_STREAM_MAX_UNITS = 12
DISPLAY_STREAM_BACKLOG_DIVISOR = 24
SCROLL_FRAME_INTERVAL_MS = 8
ACTIVE_STATUS_MIN_VISIBLE_MS = 900
SCROLL_ANIMATION_FACTOR = 0.24
SCROLL_MIN_STEP = 1.0
MAX_MATH_CACHE_ITEMS = 128
DEFAULT_BLOCK_WIDTH = 560
INLINE_MATH_HEIGHT_FACTOR = 1.04
INLINE_TALL_MATH_HEIGHT_FACTOR = 1.32
BLOCK_MATH_TARGET_HEIGHT_FACTOR = 2.6
BLOCK_MATH_MAX_SCALE = 2.8
BLOCK_MATH_SCALE_MULTIPLIER = 1.45
MATRIX_BLOCK_MATH_SCALE_MULTIPLIER = 2.0
BLOCK_MATH_MIN_VERTICAL_PADDING_PX = 10
BLOCK_MATH_VERTICAL_PADDING_FACTOR = 0.5
INLINE_MATH_RENDER_BORDER_PT = 1
BLOCK_MATH_RENDER_BORDER_PT = 5
MATRIX_BLOCK_MATH_RENDER_BORDER_PT = 8
MATH_RENDER_CACHE_VERSION = "v2"
AUTO_SCROLL_BOTTOM_THRESHOLD = 24
MAX_ATTACHMENT_TEXT_CHARS = 60000
SUPPORTED_DOCUMENT_EXTENSIONS = {".md", ".txt", ".docx", ".doc", ".xlsx", ".xls", ".csv", ".html"}
SUPPORTED_IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".gif"}
HEADING_FONT_PIXELS = {
    1: 24,
    2: 21,
    3: 18,
}
INLINE_MARKDOWN_TRANSLATION = str.maketrans({
    "＊": "*",
    "∗": "*",
    "﹡": "*",
    "✱": "*",
    "✲": "*",
    "✳": "*",
    "＿": "_",
})


@dataclass(frozen=True)
class GuiText:
    placeholder: str
    welcome: str
    pending: str
    copy_message_tooltip: str
    copied_message_tooltip: str
    new_chat_tooltip: str
    remove_image: str
    selected_image: str
    selected_file: str
    choose_image_title: str
    unsupported_file: str
    file_read_failed: str
    image_fallback: str
    file_fallback: str
    file_prompt_intro: str
    file_prompt_header: str
    file_prompt_truncated: str
    retrieving_context: str
    loading_model: str
    generating_answer: str
    model_loaded: str
    prefill_estimate: str
    compressing_context: str
    context_cleared: str
    generation_failed: str
    error_status: str
    error_title: str
    completed_status: str


GUI_TEXT_BY_LOCALE = {
    "zh": GuiText(
        placeholder="输入消息...",
        welcome="你好！我是你的本地 AI 助手，有什么我可以帮你的吗？",
        pending="正在生成...",
        copy_message_tooltip="复制消息",
        copied_message_tooltip="已复制",
        new_chat_tooltip="新建对话",
        remove_image="移除",
        selected_image="已选择图片：{name}",
        selected_file="已选择文件：{name}",
        choose_image_title="选择文件或图片",
        unsupported_file="不支持的文件类型。仅支持 md、txt、docx、doc、xlsx、xls、csv、html。",
        file_read_failed="无法读取文件：{error}",
        image_fallback="[图片]",
        file_fallback="[文件：{name}]",
        file_prompt_intro="请阅读这个文件。",
        file_prompt_header="[上传文件：{name}]\n文件内容：",
        file_prompt_truncated="\n\n[文件内容过长，已截断。]",
        retrieving_context="正在检索上下文...",
        loading_model="正在加载模型...",
        generating_answer="正在生成回答...",
        model_loaded="模型已加载",
        prefill_estimate="正在预填充上下文... 预计 {seconds:.1f}s",
        compressing_context="正在压缩上下文...",
        context_cleared="上下文已清空",
        generation_failed="生成失败。",
        error_status="发生错误",
        error_title="错误",
        completed_status="完成 · 提示 {prompt} · 生成 {generation} · 峰值 {peak:.2f} GB",
    ),
    "en": GuiText(
        placeholder="Message...",
        welcome="Hello! I am your local AI assistant. How can I help?",
        pending="Generating...",
        copy_message_tooltip="Copy message",
        copied_message_tooltip="Copied",
        new_chat_tooltip="New chat",
        remove_image="Remove",
        selected_image="Selected image: {name}",
        selected_file="Selected file: {name}",
        choose_image_title="Choose file or image",
        unsupported_file="Unsupported file type. Supported file types: md, txt, docx, doc, xlsx, xls, csv, html.",
        file_read_failed="Could not read file: {error}",
        image_fallback="[image]",
        file_fallback="[file: {name}]",
        file_prompt_intro="Please read this file.",
        file_prompt_header="[Uploaded file: {name}]\nFile content:",
        file_prompt_truncated="\n\n[File content was too long and has been truncated.]",
        retrieving_context="Retrieving context...",
        loading_model="Loading model...",
        generating_answer="Generating answer...",
        model_loaded="Model loaded",
        prefill_estimate="Prefilling context... Estimated {seconds:.1f}s",
        compressing_context="Compressing context...",
        context_cleared="Context cleared",
        generation_failed="Generation failed.",
        error_status="Error occurred",
        error_title="Error",
        completed_status="Done · prompt {prompt} · generation {generation} · peak {peak:.2f} GB",
    ),
}


def _load_local_env_values() -> dict[str, str]:
    env_path = Path(__file__).with_name(".env")
    values: dict[str, str] = {}
    try:
        lines = env_path.read_text(encoding="utf-8").splitlines()
    except OSError:
        return values
    for raw_line in lines:
        line = raw_line.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        key, value = line.split("=", 1)
        key = key.strip()
        value = value.strip()
        if len(value) >= 2 and value[0] == value[-1] and value[0] in {"'", '"'}:
            value = value[1:-1]
        if key:
            values[key] = value
    return values


def _configured_locale() -> str:
    value = os.getenv("LLM_LOCALE") or _load_local_env_values().get("LLM_LOCALE") or "zh-CN"
    normalized = value.strip().lower()
    if normalized.startswith("en"):
        return "en"
    return "zh"


GUI_TEXT = GUI_TEXT_BY_LOCALE[_configured_locale()]
PLACEHOLDER = GUI_TEXT.placeholder
WELCOME_TEXT = GUI_TEXT.welcome
PENDING_TEXT = GUI_TEXT.pending


class AttachmentReadError(RuntimeError):
    pass


@dataclass(frozen=True)
class FileAttachment:
    path: Path
    text: str
    truncated: bool = False

    @property
    def name(self) -> str:
        return self.path.name


class _HTMLTextExtractor(HTMLParser):
    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self._parts: list[str] = []

    def handle_starttag(self, tag: str, attrs) -> None:  # type: ignore[override]
        del attrs
        if tag.lower() in {"br", "p", "div", "section", "article", "tr", "li", "h1", "h2", "h3", "h4", "h5", "h6"}:
            self._parts.append("\n")

    def handle_data(self, data: str) -> None:  # type: ignore[override]
        if data.strip():
            self._parts.append(data)

    def text(self) -> str:
        return _normalize_extracted_text(" ".join(self._parts))


def _normalize_extracted_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t\f\v]+", " ", text)
    text = re.sub(r" *\n *", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _truncate_attachment_text(text: str) -> tuple[str, bool]:
    if len(text) <= MAX_ATTACHMENT_TEXT_CHARS:
        return text, False
    return text[:MAX_ATTACHMENT_TEXT_CHARS].rstrip(), True


def _read_text_path(path: Path) -> str:
    data = path.read_bytes()
    for encoding in ("utf-8-sig", "utf-16", "gb18030", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _read_html_path(path: Path) -> str:
    parser = _HTMLTextExtractor()
    parser.feed(_read_text_path(path))
    parser.close()
    return parser.text()


def _read_csv_path(path: Path) -> str:
    text = _read_text_path(path)
    rows: list[str] = []
    try:
        reader = csv.reader(io.StringIO(text))
        for row in reader:
            rows.append("\t".join(cell.strip() for cell in row))
    except csv.Error:
        return text
    return "\n".join(rows)


def _xml_text(element) -> str:
    parts: list[str] = []
    for child in element.iter():
        tag = child.tag.rsplit("}", 1)[-1]
        if tag in {"t", "v"} and child.text:
            parts.append(child.text)
    return "".join(parts)


def _read_docx_zip_path(path: Path) -> str:
    try:
        with zipfile.ZipFile(path) as archive:
            xml_data = archive.read("word/document.xml")
    except (KeyError, zipfile.BadZipFile) as exc:
        raise AttachmentReadError("Could not read the .docx document body.") from exc

    root = ElementTree.fromstring(xml_data)
    parts: list[str] = []
    for block in root.iter():
        tag = block.tag.rsplit("}", 1)[-1]
        if tag == "p":
            text = "".join(child.text or "" for child in block.iter() if child.tag.rsplit("}", 1)[-1] == "t").strip()
            if text:
                parts.append(text)
        elif tag == "tr":
            cells = []
            for cell in block:
                if cell.tag.rsplit("}", 1)[-1] == "tc":
                    cells.append(_xml_text(cell).strip())
            if any(cells):
                parts.append("\t".join(cells))
    return "\n".join(parts)


def _read_docx_path(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        return _read_docx_zip_path(path)

    document = Document(str(path))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append("\t".join(cells))
    return "\n".join(parts)


def _extract_binary_strings(path: Path) -> str:
    data = path.read_bytes()
    strings: list[str] = []
    for match in re.findall(rb"[\x20-\x7e]{4,}", data):
        text = match.decode("latin-1", errors="ignore").strip()
        if text:
            strings.append(text)
    for match in re.findall(rb"(?:[\x20-\x7e]\x00){4,}", data):
        text = match.decode("utf-16le", errors="ignore").strip()
        if text:
            strings.append(text)
    unique: list[str] = []
    seen: set[str] = set()
    for text in strings:
        if text not in seen:
            seen.add(text)
            unique.append(text)
    return "\n".join(unique)


def _read_doc_path(path: Path) -> str:
    textutil = shutil.which("textutil")
    if textutil is None:
        raise AttachmentReadError("macOS textutil is required to read .doc files.")
    result = subprocess.run(
        [textutil, "-convert", "txt", "-stdout", str(path)],
        check=False,
        capture_output=True,
        text=True,
        timeout=30,
    )
    if result.returncode != 0:
        message = result.stderr.strip() or "textutil failed to convert the .doc file."
        raise AttachmentReadError(message)
    return result.stdout


def _read_xlsx_path(path: Path) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        return _read_xlsx_zip_path(path)

    workbook = load_workbook(str(path), read_only=True, data_only=True)
    parts: list[str] = []
    try:
        for sheet in workbook.worksheets:
            parts.append(f"# {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                cells = ["" if value is None else str(value) for value in row]
                if any(cell.strip() for cell in cells):
                    parts.append("\t".join(cells).rstrip())
    finally:
        workbook.close()
    return "\n".join(parts)


def _read_xlsx_zip_path(path: Path) -> str:
    try:
        archive = zipfile.ZipFile(path)
    except zipfile.BadZipFile as exc:
        raise AttachmentReadError("Could not read the .xlsx workbook.") from exc

    with archive:
        shared_strings: list[str] = []
        try:
            shared_root = ElementTree.fromstring(archive.read("xl/sharedStrings.xml"))
        except KeyError:
            shared_root = None
        if shared_root is not None:
            for item in shared_root:
                shared_strings.append(_xml_text(item))

        sheet_paths = sorted(
            name
            for name in archive.namelist()
            if name.startswith("xl/worksheets/sheet") and name.endswith(".xml")
        )
        parts: list[str] = []
        for index, sheet_path in enumerate(sheet_paths, start=1):
            parts.append(f"# Sheet {index}")
            root = ElementTree.fromstring(archive.read(sheet_path))
            for row in root.iter():
                if row.tag.rsplit("}", 1)[-1] != "row":
                    continue
                cells: list[str] = []
                for cell in row:
                    if cell.tag.rsplit("}", 1)[-1] != "c":
                        continue
                    cell_type = cell.attrib.get("t")
                    raw_value = ""
                    if cell_type == "inlineStr":
                        raw_value = _xml_text(cell)
                    else:
                        value_node = next(
                            (child for child in cell if child.tag.rsplit("}", 1)[-1] == "v"),
                            None,
                        )
                        raw_value = "" if value_node is None or value_node.text is None else value_node.text
                        if cell_type == "s":
                            try:
                                raw_value = shared_strings[int(raw_value)]
                            except (ValueError, IndexError):
                                pass
                    cells.append(str(raw_value).strip())
                if any(cells):
                    parts.append("\t".join(cells).rstrip())
        return "\n".join(parts)


def _read_xls_path(path: Path) -> str:
    try:
        import xlrd
    except ImportError as exc:
        text = _extract_binary_strings(path)
        if text:
            return text
        raise AttachmentReadError("xlrd is required to read this .xls file.") from exc

    workbook = xlrd.open_workbook(str(path), on_demand=True)
    parts: list[str] = []
    try:
        for sheet in workbook.sheets():
            parts.append(f"# {sheet.name}")
            for row_index in range(sheet.nrows):
                cells = [str(sheet.cell_value(row_index, col_index)).strip() for col_index in range(sheet.ncols)]
                if any(cells):
                    parts.append("\t".join(cells).rstrip())
    finally:
        workbook.release_resources()
    return "\n".join(parts)


def read_file_attachment(path: str | Path) -> FileAttachment:
    file_path = Path(path)
    extension = file_path.suffix.lower()
    if extension not in SUPPORTED_DOCUMENT_EXTENSIONS:
        raise AttachmentReadError(GUI_TEXT.unsupported_file)

    if extension in {".md", ".txt"}:
        text = _read_text_path(file_path)
    elif extension == ".csv":
        text = _read_csv_path(file_path)
    elif extension == ".html":
        text = _read_html_path(file_path)
    elif extension == ".docx":
        text = _read_docx_path(file_path)
    elif extension == ".doc":
        text = _read_doc_path(file_path)
    elif extension == ".xlsx":
        text = _read_xlsx_path(file_path)
    elif extension == ".xls":
        text = _read_xls_path(file_path)
    else:
        raise AttachmentReadError(GUI_TEXT.unsupported_file)

    text = _normalize_extracted_text(text)
    if not text:
        raise AttachmentReadError("No readable text was found in the file.")
    text, truncated = _truncate_attachment_text(text)
    return FileAttachment(path=file_path, text=text, truncated=truncated)


@dataclass(frozen=True)
class TextBlock:
    lines: list[MarkdownLine]


@dataclass(frozen=True)
class CodeBlock:
    text: str


@dataclass(frozen=True)
class MathBlock:
    latex: str
    raw_source: str


@dataclass(frozen=True)
class MarkdownLine:
    kind: str
    text: str
    marker: str = ""
    indent: int = 0
    level: int = 0


@dataclass(frozen=True)
class ThematicBreak:
    pass


@dataclass(frozen=True)
class TableBlock:
    headers: list[str]
    aligns: list[str]
    rows: list[list[str]]


@dataclass(frozen=True)
class TextRun:
    text: str
    bold: bool = False


@dataclass(frozen=True)
class CodeRun:
    text: str


@dataclass(frozen=True)
class MathRun:
    latex: str
    raw_source: str


@dataclass(frozen=True)
class MathRenderArtifact:
    format: str
    data: bytes


@dataclass(frozen=True)
class MathRenderTicket:
    owner_id: str
    revision: int
    segment_id: str
    cache_key: str
    latex: str
    display: str


@dataclass(frozen=True)
class MathRenderEvent:
    owner_id: str
    revision: int
    segment_id: str
    cache_key: str
    artifact: MathRenderArtifact | None
    error: str | None = None


@dataclass(frozen=True)
class BlockState:
    block_id: str
    block: TextBlock | CodeBlock | TableBlock | MathBlock | ThematicBreak
    revision: int = 0
    stable: bool = True


@dataclass(frozen=True)
class RenderPatch:
    action: str
    block_state: BlockState | None = None
    index: int | None = None


@dataclass(frozen=True)
class ParsedStreamingBlocks:
    blocks: list[TextBlock | CodeBlock | TableBlock | MathBlock | ThematicBreak]
    stable_count: int


class FlowLayout(QLayout):
    def __init__(self, parent: QWidget | None = None, margin: int = 0, spacing: int = 4) -> None:
        super().__init__(parent)
        self._items: list[QLayoutItem] = []
        self.setContentsMargins(margin, margin, margin, margin)
        self.setSpacing(spacing)

    def addItem(self, item: QLayoutItem) -> None:
        self._items.append(item)

    def count(self) -> int:
        return len(self._items)

    def itemAt(self, index: int) -> QLayoutItem | None:
        if 0 <= index < len(self._items):
            return self._items[index]
        return None

    def takeAt(self, index: int) -> QLayoutItem | None:
        if 0 <= index < len(self._items):
            return self._items.pop(index)
        return None

    def expandingDirections(self) -> Qt.Orientations:
        return Qt.Orientations()

    def hasHeightForWidth(self) -> bool:
        return True

    def heightForWidth(self, width: int) -> int:
        return self._do_layout(QRect(0, 0, max(0, width), 0), True)

    def setGeometry(self, rect: QRect) -> None:
        super().setGeometry(rect)
        self._do_layout(rect, False)

    def sizeHint(self) -> QSize:
        return self.minimumSize()

    def minimumSize(self) -> QSize:
        size = QSize()
        for item in self._items:
            size = size.expandedTo(item.minimumSize())
        margins = self.contentsMargins()
        size += QSize(margins.left() + margins.right(), margins.top() + margins.bottom())
        return size

    def _do_layout(self, rect: QRect, test_only: bool) -> int:
        margins = self.contentsMargins()
        area = rect.adjusted(margins.left(), margins.top(), -margins.right(), -margins.bottom())
        x = area.x()
        y = area.y()
        line_height = 0
        spacing = self.spacing()
        max_right = max(area.x(), area.right())

        for item in self._items:
            hint = item.sizeHint()
            next_x = x + hint.width()
            if line_height > 0 and next_x > max_right + 1:
                x = area.x()
                y += line_height + spacing
                next_x = x + hint.width()
                line_height = 0
            if not test_only:
                item.setGeometry(QRect(QPoint(x, y), hint))
            x = next_x + spacing
            line_height = max(line_height, hint.height())

        total_height = y + line_height - rect.y() + margins.bottom()
        return max(total_height, margins.top() + margins.bottom())


class TextTokenWidget(QLabel):
    def __init__(self, text: str, object_name: str = "assistantTextToken", bold: bool = False) -> None:
        super().__init__(text)
        self.setObjectName(object_name)
        self.setProperty("markdownBold", bold)
        self.setTextFormat(Qt.PlainText)
        self.setWordWrap(False)
        self.setSizePolicy(QSizePolicy.Maximum, QSizePolicy.Fixed)
        self.setTextInteractionFlags(Qt.TextSelectableByMouse)


class InlineCodeTokenWidget(TextTokenWidget):
    def __init__(self, text: str) -> None:
        super().__init__(text, "assistantCodeToken")


def _is_matrix_like_latex(source: str) -> bool:
    return bool(
        re.search(
            r"\\begin\s*\{\s*(?:[pbBvV]?matrix|smallmatrix|array)\s*\}",
            source,
        )
        or re.search(r"\\matrix\s*\{", source)
    )


def _block_math_vertical_padding(font_height: int) -> int:
    return max(
        BLOCK_MATH_MIN_VERTICAL_PADDING_PX,
        int(round(max(1, font_height) * BLOCK_MATH_VERTICAL_PADDING_FACTOR)),
    )


class SvgArtifactWidget(QWidget):
    def __init__(
        self,
        data: bytes,
        display: str,
        inline_height_factor: float = INLINE_MATH_HEIGHT_FACTOR,
        block_scale_multiplier: float = 1.0,
    ) -> None:
        super().__init__()
        self._display = display
        self._inline_height_factor = inline_height_factor
        self._block_scale_multiplier = max(1.0, block_scale_multiplier)
        self._renderer = QSvgRenderer(QByteArray(data), self) if QSvgRenderer is not None else None
        default_size = self._renderer.defaultSize() if self._renderer is not None else QSize()
        if not default_size.isValid() or default_size.width() <= 0 or default_size.height() <= 0:
            default_size = QSize(120, 40 if display == "inline" else 72)
        self._default_size = default_size
        self.setSizePolicy(
            QSizePolicy.Maximum if display == "inline" else QSizePolicy.Preferred,
            QSizePolicy.Fixed,
        )

    def hasHeightForWidth(self) -> bool:
        return self._display == "block"

    def heightForWidth(self, width: int) -> int:
        return self._scaled_size(width).height()

    def sizeHint(self) -> QSize:
        width = self.width() if self.width() > 0 else DEFAULT_BLOCK_WIDTH
        return self._scaled_size(width)

    def minimumSizeHint(self) -> QSize:
        return self.sizeHint()

    def _scaled_size(self, width_limit: int) -> QSize:
        content_size = self._scaled_content_size(width_limit)
        if self._display == "inline":
            return content_size
        vertical_padding = _block_math_vertical_padding(self.fontMetrics().height())
        return QSize(content_size.width(), content_size.height() + vertical_padding * 2)

    def _scaled_content_size(self, width_limit: int) -> QSize:
        width_limit = max(1, width_limit)
        source = self._default_size
        if self._display == "inline":
            target_height = max(16, int(self.fontMetrics().height() * self._inline_height_factor))
            scale = target_height / max(1, source.height())
            return QSize(max(1, int(source.width() * scale)), target_height)
        width_scale = width_limit / max(1, source.width())
        desired_height = max(40, int(self.fontMetrics().height() * BLOCK_MATH_TARGET_HEIGHT_FACTOR))
        desired_scale = max(1.0, desired_height / max(1, source.height()))
        base_scale = min(BLOCK_MATH_MAX_SCALE, width_scale, desired_scale)
        scale = min(
            BLOCK_MATH_MAX_SCALE * self._block_scale_multiplier,
            width_scale,
            base_scale * self._block_scale_multiplier,
        )
        return QSize(max(1, int(source.width() * scale)), max(1, int(source.height() * scale)))

    def _target_rect(self) -> QRectF:
        target_size = self._scaled_content_size(self.width() if self.width() > 0 else DEFAULT_BLOCK_WIDTH)
        x = 0.0
        if self._display == "block":
            x = max(0.0, (self.width() - target_size.width()) / 2.0)
        y = max(0.0, (self.height() - target_size.height()) / 2.0)
        return QRectF(x, y, float(target_size.width()), float(target_size.height()))

    def paintEvent(self, event) -> None:  # type: ignore[override]
        del event
        if self._renderer is None:
            return
        painter = QPainter(self)
        painter.setRenderHint(QPainter.Antialiasing)
        self._renderer.render(painter, self._target_rect())


class RasterArtifactWidget(QWidget):
    def __init__(
        self,
        data: bytes,
        display: str,
        inline_height_factor: float = INLINE_MATH_HEIGHT_FACTOR,
        block_scale_multiplier: float = 1.0,
    ) -> None:
        super().__init__()
        self._display = display
        self._inline_height_factor = inline_height_factor
        self._block_scale_multiplier = max(1.0, block_scale_multiplier)
        self._pixmap = QPixmap()
        self._pixmap.loadFromData(data)
        if self._pixmap.isNull():
            self._pixmap = QPixmap(120, 40 if display == "inline" else 72)
            self._pixmap.fill(Qt.transparent)
        self.setSizePolicy(
            QSizePolicy.Maximum if display == "inline" else QSizePolicy.Preferred,
            QSizePolicy.Fixed,
        )

    def hasHeightForWidth(self) -> bool:
        return self._display == "block"

    def heightForWidth(self, width: int) -> int:
        return self._scaled_size(width).height()

    def sizeHint(self) -> QSize:
        width = self.width() if self.width() > 0 else DEFAULT_BLOCK_WIDTH
        return self._scaled_size(width)

    def minimumSizeHint(self) -> QSize:
        return self.sizeHint()

    def _scaled_size(self, width_limit: int) -> QSize:
        content_size = self._scaled_content_size(width_limit)
        if self._display == "inline":
            return content_size
        vertical_padding = _block_math_vertical_padding(self.fontMetrics().height())
        return QSize(content_size.width(), content_size.height() + vertical_padding * 2)

    def _scaled_content_size(self, width_limit: int) -> QSize:
        width_limit = max(1, width_limit)
        source = self._pixmap.size()
        if source.width() <= 0 or source.height() <= 0:
            source = QSize(120, 40 if self._display == "inline" else 72)
        if self._display == "inline":
            target_height = max(16, int(self.fontMetrics().height() * self._inline_height_factor))
            scale = target_height / max(1, source.height())
            return QSize(max(1, int(source.width() * scale)), target_height)
        width_scale = width_limit / max(1, source.width())
        desired_height = max(40, int(self.fontMetrics().height() * BLOCK_MATH_TARGET_HEIGHT_FACTOR))
        desired_scale = max(1.0, desired_height / max(1, source.height()))
        base_scale = min(BLOCK_MATH_MAX_SCALE, width_scale, desired_scale)
        scale = min(
            BLOCK_MATH_MAX_SCALE * self._block_scale_multiplier,
            width_scale,
            base_scale * self._block_scale_multiplier,
        )
        return QSize(max(1, int(source.width() * scale)), max(1, int(source.height() * scale)))

    def _target_rect(self) -> QRect:
        target_size = self._scaled_content_size(self.width() if self.width() > 0 else DEFAULT_BLOCK_WIDTH)
        x = 0
        if self._display == "block":
            x = max(0, (self.width() - target_size.width()) // 2)
        y = max(0, (self.height() - target_size.height()) // 2)
        return QRect(x, y, target_size.width(), target_size.height())

    def paintEvent(self, event) -> None:  # type: ignore[override]
        del event
        if self._pixmap.isNull():
            return
        painter = QPainter(self)
        painter.setRenderHint(QPainter.SmoothPixmapTransform)
        painter.drawPixmap(self._target_rect(), self._pixmap)


class MathDisplayWidget(QWidget):
    def __init__(
        self,
        raw_source: str,
        display: str,
        inline_height_factor: float = INLINE_MATH_HEIGHT_FACTOR,
    ) -> None:
        super().__init__()
        self.display = display
        self._inline_height_factor = inline_height_factor
        self._block_scale_multiplier = (
            MATRIX_BLOCK_MATH_SCALE_MULTIPLIER
            if display == "block" and _is_matrix_like_latex(raw_source)
            else BLOCK_MATH_SCALE_MULTIPLIER
            if display == "block"
            else 1.0
        )
        self._artifact_widget: QWidget | None = None
        self._stack = QStackedLayout(self)
        self._stack.setContentsMargins(0, 0, 0, 0)
        self._stack.setStackingMode(QStackedLayout.StackOne)

        self._placeholder = QLabel(raw_source)
        self._placeholder.setTextFormat(Qt.PlainText)
        self._placeholder.setTextInteractionFlags(Qt.TextSelectableByMouse)
        self._placeholder.setObjectName(
            "assistantMathFallbackBlock" if display == "block" else "assistantMathFallbackInline"
        )
        self._placeholder.setWordWrap(display == "block")
        self._placeholder.setAlignment(Qt.AlignLeft | Qt.AlignVCenter)
        self._placeholder.setSizePolicy(
            QSizePolicy.Preferred if display == "block" else QSizePolicy.Maximum,
            QSizePolicy.Fixed,
        )
        self._stack.addWidget(self._placeholder)
        self._stack.setCurrentWidget(self._placeholder)
        self.setContentsMargins(0, 0, 0, 0)
        if display == "block":
            self.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
            self.setMinimumHeight(max(30, self.fontMetrics().height() * 2))
        else:
            self.setSizePolicy(QSizePolicy.Maximum, QSizePolicy.Fixed)
            self.setMinimumHeight(max(18, int(self.fontMetrics().height() * max(1.1, self._inline_height_factor))))

    def hasHeightForWidth(self) -> bool:
        current = self._stack.currentWidget()
        return bool(current and current.hasHeightForWidth())

    def heightForWidth(self, width: int) -> int:
        current = self._stack.currentWidget()
        if current is not None and current.hasHeightForWidth():
            return current.heightForWidth(width)
        if current is not None:
            return current.sizeHint().height()
        return super().heightForWidth(width)

    def sizeHint(self) -> QSize:
        current = self._stack.currentWidget()
        if current is not None:
            return current.sizeHint()
        return super().sizeHint()

    def minimumSizeHint(self) -> QSize:
        return self.sizeHint()

    def set_artifact(self, artifact: MathRenderArtifact | None) -> None:
        if artifact is None:
            self._stack.setCurrentWidget(self._placeholder)
            self.updateGeometry()
            return
        if artifact.format == "svg" and QSvgRenderer is not None:
            widget = SvgArtifactWidget(
                artifact.data,
                self.display,
                inline_height_factor=self._inline_height_factor,
                block_scale_multiplier=self._block_scale_multiplier,
            )
        else:
            widget = RasterArtifactWidget(
                artifact.data,
                self.display,
                inline_height_factor=self._inline_height_factor,
                block_scale_multiplier=self._block_scale_multiplier,
            )
        if self._artifact_widget is not None:
            self._stack.removeWidget(self._artifact_widget)
            self._artifact_widget.deleteLater()
        self._artifact_widget = widget
        self._stack.addWidget(widget)
        self._stack.setCurrentWidget(widget)
        self.updateGeometry()


class InlineFlowWidget(QWidget):
    def __init__(
        self,
        runs: list[TextRun | CodeRun | MathRun],
        owner_id: str,
        revision: int,
        next_segment_id,
        math_renderer: "MathRenderService",
        object_name: str | None = None,
        heading_level: int = 0,
    ) -> None:
        super().__init__()
        if object_name:
            self.setObjectName(object_name)
        if heading_level > 0:
            self.setProperty("headingLevel", heading_level)
            heading_font = self.font()
            heading_font.setPixelSize(_heading_font_pixel_size(heading_level))
            heading_font.setWeight(QFont.Bold)
            self.setFont(heading_font)
        self.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
        layout = FlowLayout(self, spacing=4)
        layout.setContentsMargins(0, 0, 0, 0)

        for run in runs:
            if isinstance(run, TextRun):
                for token in _tokenize_plain_text(run.text):
                    if token:
                        layout.addWidget(TextTokenWidget(token, bold=run.bold))
            elif isinstance(run, CodeRun):
                if run.text:
                    layout.addWidget(InlineCodeTokenWidget(run.text))
            else:
                segment_id = next_segment_id()
                widget = MathDisplayWidget(
                    run.raw_source,
                    "inline",
                    inline_height_factor=_inline_math_height_factor(run.latex),
                )
                layout.addWidget(widget)
                math_renderer.attach_widget(owner_id, revision, segment_id, widget)
                math_renderer.request_render(owner_id, revision, segment_id, run.latex, "inline")

    def hasHeightForWidth(self) -> bool:
        return True

    def heightForWidth(self, width: int) -> int:
        layout = self.layout()
        if layout is not None and layout.hasHeightForWidth():
            return layout.heightForWidth(width)
        return self.sizeHint().height()


class MarkdownLineWidget(QWidget):
    def __init__(
        self,
        line: MarkdownLine,
        owner_id: str,
        revision: int,
        next_segment_id,
        math_renderer: "MathRenderService",
    ) -> None:
        super().__init__()
        self.line = line
        self.setObjectName(f"assistantLine_{line.kind}")
        self.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)

        if line.kind == "heading":
            layout = QHBoxLayout(self)
            layout.setContentsMargins(0, 2, 0, 3)
            layout.setSpacing(0)
            flow = InlineFlowWidget(
                _parse_inline_runs(line.text),
                owner_id,
                revision,
                next_segment_id,
                math_renderer,
                object_name="assistantHeadingFlow",
                heading_level=line.level,
            )
            layout.addWidget(flow, 1)
            return

        outer = QHBoxLayout(self)
        left_padding = line.indent * 10
        outer.setContentsMargins(left_padding, 0, 0, 0)
        outer.setSpacing(6)

        marker_text = ""
        marker_name = None
        if line.kind == "bullet":
            marker_text = "•"
            marker_name = "assistantListMarker"
        elif line.kind == "ordered":
            marker_text = line.marker
            marker_name = "assistantListMarker"
        elif line.kind == "blockquote":
            marker_text = "│"
            marker_name = "assistantQuoteMarker"

        if marker_name is not None:
            marker = QLabel(marker_text)
            marker.setObjectName(marker_name)
            marker.setAlignment(Qt.AlignLeft | Qt.AlignTop)
            marker.setFixedWidth(22 if line.kind == "ordered" else 16)
            outer.addWidget(marker, 0, Qt.AlignTop)

        flow = InlineFlowWidget(_parse_inline_runs(line.text), owner_id, revision, next_segment_id, math_renderer)
        outer.addWidget(flow, 1)

    def hasHeightForWidth(self) -> bool:
        layout = self.layout()
        return bool(layout and layout.hasHeightForWidth())

    def heightForWidth(self, width: int) -> int:
        layout = self.layout()
        if layout is not None and layout.hasHeightForWidth():
            return layout.heightForWidth(width)
        return self.sizeHint().height()


class ThematicBreakWidget(QWidget):
    def __init__(self) -> None:
        super().__init__()
        self.setFixedHeight(12)
        self.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)

    def paintEvent(self, event) -> None:  # type: ignore[override]
        del event
        painter = QPainter(self)
        painter.setRenderHint(QPainter.Antialiasing)
        pen = QPen(QColor(0, 0, 0, 35))
        pen.setWidth(1)
        painter.setPen(pen)
        y = self.rect().center().y()
        painter.drawLine(0, y, self.width(), y)


class CodeBlockWidget(QFrame):
    def __init__(self, text: str) -> None:
        super().__init__()
        self.setObjectName("assistantCodeBlock")
        self.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
        layout = QVBoxLayout(self)
        layout.setContentsMargins(10, 8, 10, 8)
        layout.setSpacing(0)

        label = QLabel(text)
        label.setObjectName("assistantCodeBlockLabel")
        label.setTextFormat(Qt.PlainText)
        label.setWordWrap(True)
        label.setTextInteractionFlags(Qt.TextSelectableByMouse)
        label.setAlignment(Qt.AlignLeft | Qt.AlignTop)
        layout.addWidget(label)


class MarkdownTableCellWidget(QFrame):
    def __init__(
        self,
        text: str,
        alignment: str,
        table_header: bool,
        alternate_row: bool,
        owner_id: str,
        revision: int,
        next_segment_id,
        math_renderer: "MathRenderService",
    ) -> None:
        super().__init__()
        self.setObjectName("assistantTableCell")
        self.setProperty("tableHeader", table_header)
        self.setProperty("alternateRow", alternate_row and not table_header)
        self.setProperty("tableAlign", alignment)
        self.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Preferred)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(10, 8, 10, 8)
        layout.setSpacing(0)

        runs = _parse_inline_runs(text)
        if table_header:
            runs = _promote_runs_to_table_header(runs)

        if runs:
            layout.addWidget(InlineFlowWidget(runs, owner_id, revision, next_segment_id, math_renderer))
        else:
            spacer = QWidget()
            spacer.setFixedHeight(max(18, self.fontMetrics().height()))
            layout.addWidget(spacer)

    def hasHeightForWidth(self) -> bool:
        layout = self.layout()
        return bool(layout and layout.hasHeightForWidth())

    def heightForWidth(self, width: int) -> int:
        layout = self.layout()
        if layout is not None and layout.hasHeightForWidth():
            return layout.heightForWidth(width)
        return self.sizeHint().height()


class MarkdownTableRowWidget(QWidget):
    def __init__(
        self,
        cells: list[str],
        aligns: list[str],
        table_header: bool,
        alternate_row: bool,
        owner_id: str,
        revision: int,
        next_segment_id,
        math_renderer: "MathRenderService",
    ) -> None:
        super().__init__()
        self.setObjectName("assistantTableRow")
        self.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)

        layout = QHBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(1)

        for index, text in enumerate(cells):
            alignment = aligns[index] if index < len(aligns) else "left"
            layout.addWidget(
                MarkdownTableCellWidget(
                    text=text,
                    alignment=alignment,
                    table_header=table_header,
                    alternate_row=alternate_row,
                    owner_id=owner_id,
                    revision=revision,
                    next_segment_id=next_segment_id,
                    math_renderer=math_renderer,
                ),
                1,
            )

    def hasHeightForWidth(self) -> bool:
        layout = self.layout()
        return bool(layout and layout.hasHeightForWidth())

    def heightForWidth(self, width: int) -> int:
        layout = self.layout()
        if layout is not None and layout.hasHeightForWidth():
            return layout.heightForWidth(width)
        return self.sizeHint().height()


class MarkdownTableWidget(QFrame):
    def __init__(
        self,
        table: TableBlock,
        owner_id: str,
        revision: int,
        next_segment_id,
        math_renderer: "MathRenderService",
    ) -> None:
        super().__init__()
        self.setObjectName("assistantTable")
        self.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(1, 1, 1, 1)
        layout.setSpacing(1)

        layout.addWidget(
            MarkdownTableRowWidget(
                cells=table.headers,
                aligns=table.aligns,
                table_header=True,
                alternate_row=False,
                owner_id=owner_id,
                revision=revision,
                next_segment_id=next_segment_id,
                math_renderer=math_renderer,
            )
        )

        for row_index, row in enumerate(table.rows):
            layout.addWidget(
                MarkdownTableRowWidget(
                    cells=row,
                    aligns=table.aligns,
                    table_header=False,
                    alternate_row=row_index % 2 == 1,
                    owner_id=owner_id,
                    revision=revision,
                    next_segment_id=next_segment_id,
                    math_renderer=math_renderer,
                )
            )

    def hasHeightForWidth(self) -> bool:
        layout = self.layout()
        return bool(layout and layout.hasHeightForWidth())

    def heightForWidth(self, width: int) -> int:
        layout = self.layout()
        if layout is not None and layout.hasHeightForWidth():
            return layout.heightForWidth(width)
        return self.sizeHint().height()


class BlockMathWidget(QWidget):
    def __init__(
        self,
        raw_source: str,
        latex: str,
        owner_id: str,
        revision: int,
        segment_id: str,
        math_renderer: "MathRenderService",
    ) -> None:
        super().__init__()
        self.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 6, 0, 8)
        layout.setSpacing(0)
        self.math_widget = MathDisplayWidget(raw_source, "block")
        layout.addWidget(self.math_widget)
        math_renderer.attach_widget(owner_id, revision, segment_id, self.math_widget)
        math_renderer.request_render(owner_id, revision, segment_id, latex, "block")

    def hasHeightForWidth(self) -> bool:
        layout = self.layout()
        return bool(layout and layout.hasHeightForWidth())

    def heightForWidth(self, width: int) -> int:
        layout = self.layout()
        if layout is not None and layout.hasHeightForWidth():
            return layout.heightForWidth(width)
        return self.sizeHint().height()


class AssistantContentWidget(QWidget):
    content_height_changed = Signal()

    def __init__(self, math_renderer: "MathRenderService") -> None:
        super().__init__()
        self.setObjectName("assistantContent")
        self.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
        self._math_renderer = math_renderer
        self._owner_id = f"assistant-{id(self)}"
        self._revision = 0
        self._segment_counter = 0
        self._current_text = ""
        self._stream_surface = AssistantRenderSurface(self)
        self._layout = QVBoxLayout(self)
        self._layout.setContentsMargins(0, 0, 0, 0)
        self._layout.setSpacing(6)

    def begin_stream(self) -> None:
        self._segment_counter = 0
        self._math_renderer.detach_owner(self._owner_id)
        self._clear_layout()
        self._stream_surface.begin_stream(self._current_text)

    def apply_stream_units(self, units: list[str]) -> None:
        if not units:
            return
        self._stream_surface.apply_stream_units(units)
        self._current_text = self._stream_surface.current_text

    def finish_stream(self, final_text: str) -> None:
        self._stream_surface.finish_stream(final_text)
        self._current_text = final_text

    def render_text(self, text: str) -> None:
        self._current_text = text
        self._revision += 1
        self._segment_counter = 0
        revision = self._revision
        self._math_renderer.detach_owner(self._owner_id)
        self._clear_layout()

        blocks = _parse_message_blocks(text)
        if not blocks:
            spacer = QWidget()
            spacer.setFixedHeight(1)
            self._layout.addWidget(spacer)
        else:
            for block in blocks:
                if isinstance(block, TextBlock):
                    for line in block.lines:
                        self._layout.addWidget(
                            MarkdownLineWidget(
                                line,
                                self._owner_id,
                                revision,
                                self._next_segment_id,
                                self._math_renderer,
                            )
                        )
                elif isinstance(block, CodeBlock):
                    self._layout.addWidget(CodeBlockWidget(block.text))
                elif isinstance(block, TableBlock):
                    self._layout.addWidget(
                        MarkdownTableWidget(
                            table=block,
                            owner_id=self._owner_id,
                            revision=revision,
                            next_segment_id=self._next_segment_id,
                            math_renderer=self._math_renderer,
                        )
                    )
                elif isinstance(block, MathBlock):
                    segment_id = self._next_segment_id()
                    self._layout.addWidget(
                        BlockMathWidget(
                            block.raw_source,
                            block.latex,
                            self._owner_id,
                            revision,
                            segment_id,
                            self._math_renderer,
                        )
                    )
                else:
                    self._layout.addWidget(ThematicBreakWidget())
        self.content_height_changed.emit()

    def _render_block_widget(self, block_state: BlockState) -> QWidget:
        block = block_state.block
        revision = block_state.revision
        if isinstance(block, TextBlock):
            container = QWidget()
            container.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
            layout = QVBoxLayout(container)
            layout.setContentsMargins(0, 0, 0, 0)
            layout.setSpacing(6)
            for line in block.lines:
                layout.addWidget(
                    MarkdownLineWidget(
                        line,
                        self._owner_id,
                        revision,
                        self._next_segment_id,
                        self._math_renderer,
                    )
                )
            return container
        if isinstance(block, CodeBlock):
            return CodeBlockWidget(block.text)
        if isinstance(block, TableBlock):
            return MarkdownTableWidget(
                table=block,
                owner_id=self._owner_id,
                revision=revision,
                next_segment_id=self._next_segment_id,
                math_renderer=self._math_renderer,
            )
        if isinstance(block, MathBlock):
            segment_id = self._next_segment_id()
            return BlockMathWidget(
                block.raw_source,
                block.latex,
                self._owner_id,
                revision,
                segment_id,
                self._math_renderer,
            )
        return ThematicBreakWidget()

    def _replace_layout_widget(self, index: int, widget: QWidget) -> None:
        old_item = self._layout.takeAt(index)
        if old_item is not None:
            old_widget = old_item.widget()
            if old_widget is not None:
                old_widget.setParent(None)
                old_widget.deleteLater()
        self._layout.insertWidget(index, widget)

    def _apply_render_patches(self, patches: list[RenderPatch]) -> None:
        if not patches:
            return
        for patch in patches:
            if patch.action == "remove_from":
                start = patch.index if patch.index is not None else self._layout.count()
                while self._layout.count() > start:
                    item = self._layout.takeAt(start)
                    if item is None:
                        break
                    widget = item.widget()
                    if widget is not None:
                        widget.setParent(None)
                        widget.deleteLater()
                continue
            if patch.block_state is None:
                continue
            index = patch.index if patch.index is not None else self._layout.count()
            widget = self._render_block_widget(patch.block_state)
            if patch.action == "append":
                self._layout.addWidget(widget)
            elif patch.action == "replace" and 0 <= index < self._layout.count():
                self._replace_layout_widget(index, widget)
            elif patch.action == "insert":
                self._layout.insertWidget(max(0, min(index, self._layout.count())), widget)
        self.content_height_changed.emit()

    def _next_segment_id(self) -> str:
        segment_id = f"seg-{self._segment_counter}"
        self._segment_counter += 1
        return segment_id

    def _clear_layout(self) -> None:
        while self._layout.count() > 0:
            item = self._layout.takeAt(0)
            widget = item.widget() if item is not None else None
            if widget is not None:
                widget.deleteLater()

    def hasHeightForWidth(self) -> bool:
        return True

    def heightForWidth(self, width: int) -> int:
        if self._layout.hasHeightForWidth():
            return self._layout.heightForWidth(width)
        return self._layout.sizeHint().height()

    def sizeHint(self) -> QSize:
        size = self._layout.sizeHint()
        if not size.isValid():
            size = QSize(self.width() if self.width() > 0 else DEFAULT_BLOCK_WIDTH, 1)
        return size


class StreamingDocumentState:
    def __init__(self) -> None:
        self.source_text = ""
        self.blocks: list[BlockState] = []
        self._block_ids: dict[tuple[int, str], str] = {}
        self._next_block_id = 0

    def reset(self, text: str = "") -> None:
        self.source_text = text
        self.blocks.clear()
        self._block_ids.clear()
        self._next_block_id = 0

    def append_units(self, units: list[str]) -> None:
        self.source_text += "".join(units)

    def reconcile(self, text: str, *, final: bool = False) -> list[RenderPatch]:
        self.source_text = text
        if final:
            final_blocks = _parse_message_blocks(text)
            parsed = ParsedStreamingBlocks(final_blocks, len(final_blocks))
        else:
            parsed = _parse_streaming_message_blocks(text)
        next_blocks: list[BlockState] = []
        for index, block in enumerate(parsed.blocks):
            fingerprint = _block_fingerprint(block)
            key = (index, fingerprint)
            block_id = self._block_ids.get(key)
            if block_id is None and index < len(self.blocks) and not self.blocks[index].stable:
                block_id = self.blocks[index].block_id
            if block_id is None:
                block_id = f"block-{self._next_block_id}"
                self._next_block_id += 1
            revision = _stable_int_hash(f"{block_id}:{fingerprint}")
            next_blocks.append(
                BlockState(
                    block_id=block_id,
                    block=block,
                    revision=revision,
                    stable=index < parsed.stable_count,
                )
            )

        common = 0
        for old, new in zip(self.blocks, next_blocks):
            if old.block_id != new.block_id or _block_fingerprint(old.block) != _block_fingerprint(new.block):
                break
            common += 1

        patches: list[RenderPatch] = []
        if common < len(self.blocks):
            patches.append(RenderPatch("remove_from", index=common))
        for index in range(common, len(next_blocks)):
            patches.append(RenderPatch("append", block_state=next_blocks[index], index=index))

        self.blocks = next_blocks
        self._block_ids = {(index, _block_fingerprint(state.block)): state.block_id for index, state in enumerate(self.blocks)}
        return patches


class AssistantRenderSurface:
    def __init__(self, widget: AssistantContentWidget) -> None:
        self._widget = widget
        self._state = StreamingDocumentState()
        self._active = False

    @property
    def current_text(self) -> str:
        return self._state.source_text

    def begin_stream(self, text: str = "") -> None:
        if self._active:
            return
        self._active = True
        self._state.reset(text)
        if text:
            self._widget._apply_render_patches(self._state.reconcile(text))

    def apply_stream_units(self, units: list[str]) -> None:
        if not self._active:
            self.begin_stream("")
        self._state.append_units(units)
        patches = self._state.reconcile(self._state.source_text)
        self._widget._apply_render_patches(patches)

    def finish_stream(self, final_text: str) -> None:
        self._active = False
        patches = self._state.reconcile(final_text, final=True)
        self._widget._apply_render_patches(patches)


class MathRenderService(QObject):
    render_finished = Signal(object)

    def __init__(self, parent: QObject | None = None) -> None:
        super().__init__(parent)
        self._executor = ThreadPoolExecutor(max_workers=2, thread_name_prefix="math-render")
        self._cache: OrderedDict[str, MathRenderArtifact | None] = OrderedDict()
        self._pending: dict[str, list[MathRenderTicket]] = {}
        self._widgets: dict[tuple[str, int, str], MathDisplayWidget] = {}
        self._toolchain: tuple[str, str | None, str | None] | None = None
        self.render_finished.connect(self._apply_render_result)

    def shutdown(self) -> None:
        self._executor.shutdown(wait=False, cancel_futures=True)

    def attach_widget(self, owner_id: str, revision: int, segment_id: str, widget: MathDisplayWidget) -> None:
        key = (owner_id, revision, segment_id)
        self._widgets[key] = widget
        widget.destroyed.connect(lambda *_args, stored_key=key: self._widgets.pop(stored_key, None))

    def detach_owner_revision(self, owner_id: str, revision: int) -> None:
        if revision < 0:
            return
        stale_keys = [key for key in self._widgets if key[0] == owner_id and key[1] <= revision]
        for key in stale_keys:
            self._widgets.pop(key, None)

    def detach_owner(self, owner_id: str) -> None:
        stale_keys = [key for key in self._widgets if key[0] == owner_id]
        for key in stale_keys:
            self._widgets.pop(key, None)

    def request_render(self, owner_id: str, revision: int, segment_id: str, latex: str, display: str) -> None:
        normalized = _normalize_formula_source(latex)
        cache_key = _formula_cache_key(normalized, display)
        ticket = MathRenderTicket(owner_id, revision, segment_id, cache_key, normalized, display)

        if cache_key in self._cache:
            artifact = self._cache[cache_key]
            QTimer.singleShot(0, lambda t=ticket, a=artifact: self.render_finished.emit(MathRenderEvent(t.owner_id, t.revision, t.segment_id, t.cache_key, a)))
            return

        listeners = self._pending.setdefault(cache_key, [])
        listeners.append(ticket)
        if len(listeners) > 1:
            return

        future = self._executor.submit(self._compile_formula, normalized, display)
        future.add_done_callback(lambda done, key=cache_key: self._on_future_done(key, done))

    def _on_future_done(self, cache_key: str, future: Future) -> None:
        try:
            artifact = future.result()
            error = None
        except Exception as exc:
            artifact = None
            error = str(exc)
        self._remember_cache(cache_key, artifact)
        tickets = self._pending.pop(cache_key, [])
        for ticket in tickets:
            self.render_finished.emit(
                MathRenderEvent(
                    owner_id=ticket.owner_id,
                    revision=ticket.revision,
                    segment_id=ticket.segment_id,
                    cache_key=cache_key,
                    artifact=artifact,
                    error=error,
                )
            )

    def _remember_cache(self, cache_key: str, artifact: MathRenderArtifact | None) -> None:
        self._cache[cache_key] = artifact
        self._cache.move_to_end(cache_key)
        while len(self._cache) > MAX_MATH_CACHE_ITEMS:
            self._cache.popitem(last=False)

    @Slot(object)
    def _apply_render_result(self, event: MathRenderEvent) -> None:
        widget = self._widgets.get((event.owner_id, event.revision, event.segment_id))
        if widget is None:
            return
        widget.set_artifact(event.artifact)
        parent = widget.parentWidget()
        while parent is not None:
            if isinstance(parent, AssistantContentWidget):
                parent.content_height_changed.emit()
                break
            parent = parent.parentWidget()

    def _resolve_toolchain(self) -> tuple[str, str | None, str | None]:
        if self._toolchain is not None:
            return self._toolchain
        tex_engine = shutil.which("tectonic")
        svg_converter = shutil.which("pdf2svg") or shutil.which("pdftocairo")
        png_converter = shutil.which("pdftocairo")
        self._toolchain = (tex_engine or "", svg_converter, png_converter)
        return self._toolchain

    def _compile_formula(self, latex: str, display: str) -> MathRenderArtifact | None:
        tex_engine, svg_converter, png_converter = self._resolve_toolchain()
        if not tex_engine:
            return None
        if svg_converter is None and png_converter is None:
            return None
        if _is_probably_invalid_latex(latex):
            return None

        document = _build_tex_document(latex, display)
        with tempfile.TemporaryDirectory(prefix="local-llm-math-") as temp_dir:
            workdir = Path(temp_dir)
            tex_path = workdir / "formula.tex"
            pdf_path = workdir / "formula.pdf"
            tex_path.write_text(document, encoding="utf-8")

            tex_command = [tex_engine, "--outdir", str(workdir), tex_path.name]
            tex_result = subprocess.run(
                tex_command,
                cwd=workdir,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                timeout=60,
                check=False,
                text=True,
            )
            if tex_result.returncode != 0 or not pdf_path.exists():
                return None

            if QSvgRenderer is not None and svg_converter:
                svg_path = workdir / "formula.svg"
                if Path(svg_converter).name == "pdf2svg":
                    command = [svg_converter, str(pdf_path), str(svg_path), "1"]
                else:
                    command = [svg_converter, "-svg", str(pdf_path), str(svg_path)]
                svg_result = subprocess.run(
                    command,
                    cwd=workdir,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    timeout=30,
                    check=False,
                )
                if svg_result.returncode == 0 and svg_path.exists():
                    return MathRenderArtifact("svg", svg_path.read_bytes())

            if png_converter:
                png_base = workdir / "formula"
                png_path = workdir / "formula.png"
                png_result = subprocess.run(
                    [png_converter, "-png", "-singlefile", "-transp", str(pdf_path), str(png_base)],
                    cwd=workdir,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    timeout=30,
                    check=False,
                )
                if png_result.returncode == 0 and png_path.exists():
                    return MathRenderArtifact("png", png_path.read_bytes())

        return None


def _build_tex_document(latex: str, display: str) -> str:
    if display == "block":
        math_body = f"\\[{latex}\\]"
        border_pt = MATRIX_BLOCK_MATH_RENDER_BORDER_PT if _is_matrix_like_latex(latex) else BLOCK_MATH_RENDER_BORDER_PT
    else:
        math_body = f"\\({latex}\\)"
        border_pt = INLINE_MATH_RENDER_BORDER_PT
    return rf"""
\documentclass[preview,border={border_pt}pt]{{standalone}}
\usepackage{{amsmath,amssymb,amsfonts,mathtools,bm}}
\usepackage{{xcolor}}
\begin{{document}}
{math_body}
\end{{document}}
""".strip()


def _formula_cache_key(latex: str, display: str) -> str:
    payload = f"{MATH_RENDER_CACHE_VERSION}::{display}::{latex}".encode("utf-8")
    return hashlib.sha256(payload).hexdigest()


def _normalize_formula_source(latex: str) -> str:
    return latex.strip()


def _is_probably_invalid_latex(latex: str) -> bool:
    probes = [
        r"\notacommand",
        r"\undefined",
        r"\badcommand",
    ]
    return any(token in latex for token in probes)


def _split_display_units(text: str) -> list[str]:
    units: list[str] = []
    index = 0
    while index < len(text):
        if text.startswith("```", index):
            units.append("```")
            index += 3
            continue
        if text.startswith("**", index) or text.startswith("__", index) or text.startswith("$$", index):
            units.append(text[index : index + 2])
            index += 2
            continue
        if text[index] == "\\":
            end = index + 1
            while end < len(text) and text[end].isalpha():
                end += 1
            if end > index + 1:
                units.append(text[index:end])
                index = end
                continue
        units.append(text[index])
        index += 1
    return units


def _classify_markdown_line(line: str) -> MarkdownLine | ThematicBreak:
    stripped = line.strip()
    if stripped in {"---", "***", "___"}:
        return ThematicBreak()

    indent = len(line) - len(line.lstrip(" "))
    text = line[indent:]

    heading_match = re.match(r"^(#{1,6})\s+(.*)$", text)
    if heading_match:
        marks, content = heading_match.groups()
        return MarkdownLine(kind="heading", text=content, level=len(marks), indent=indent)

    unordered_match = re.match(r"^([-*+])\s+(.*)$", text)
    if unordered_match:
        marker, content = unordered_match.groups()
        return MarkdownLine(kind="bullet", text=content, marker=marker, indent=indent)

    ordered_match = re.match(r"^(\d+\.)\s+(.*)$", text)
    if ordered_match:
        marker, content = ordered_match.groups()
        return MarkdownLine(kind="ordered", text=content, marker=marker, indent=indent)

    blockquote_match = re.match(r"^(>+)\s?(.*)$", text)
    if blockquote_match:
        marks, content = blockquote_match.groups()
        return MarkdownLine(kind="blockquote", text=content, marker=marks, indent=indent, level=len(marks))

    return MarkdownLine(kind="paragraph", text=text, indent=indent)


def _parse_single_line_block_math(line: str) -> MathBlock | None:
    stripped = line.strip()
    if not stripped.startswith("$$") or not stripped.endswith("$$"):
        return None
    if len(stripped) <= 4:
        return None
    latex = stripped[2:-2].strip()
    if not latex:
        return None
    return MathBlock(latex=latex, raw_source=stripped)


def _inline_math_height_factor(latex: str) -> float:
    compact = latex.replace(" ", "")
    tall_markers = (
        r"\frac",
        r"\dfrac",
        r"\tfrac",
        r"\sqrt",
        r"\binom",
        r"\sum",
        r"\prod",
        r"\int",
        r"\lim",
        r"\begin{matrix",
        r"\begin{pmatrix",
        r"\begin{bmatrix",
        r"\begin{vmatrix",
        r"\begin{Vmatrix",
        r"\begin{cases",
        r"\begin{aligned",
        r"\begin{array",
    )
    if any(marker in compact for marker in tall_markers):
        return INLINE_TALL_MATH_HEIGHT_FACTOR
    return INLINE_MATH_HEIGHT_FACTOR


def _heading_font_pixel_size(level: int) -> int:
    return HEADING_FONT_PIXELS.get(level, 16)


def _normalize_inline_markdown_delimiters(text: str) -> str:
    return text.translate(INLINE_MARKDOWN_TRANSLATION)


def _match_bold_delimiter(text: str, index: int) -> str | None:
    if text.startswith("**", index):
        return "**"
    if text.startswith("__", index):
        return "__"
    return None


def _promote_runs_to_table_header(runs: list[TextRun | CodeRun | MathRun]) -> list[TextRun | CodeRun | MathRun]:
    promoted: list[TextRun | CodeRun | MathRun] = []
    for run in runs:
        if isinstance(run, TextRun):
            promoted.append(TextRun(run.text, bold=True))
        else:
            promoted.append(run)
    return promoted


def _split_table_cells(line: str) -> list[str] | None:
    stripped = line.strip()
    if "|" not in stripped:
        return None

    content = stripped
    if content.startswith("|"):
        content = content[1:]
    if content.endswith("|"):
        content = content[:-1]

    cells: list[str] = []
    current: list[str] = []
    in_code = False
    index = 0

    while index < len(content):
        char = content[index]
        if char == "\\" and index + 1 < len(content) and content[index + 1] == "|":
            current.append("|")
            index += 2
            continue
        if char == "`":
            in_code = not in_code
            current.append(char)
            index += 1
            continue
        if char == "|" and not in_code:
            cells.append("".join(current).strip())
            current.clear()
            index += 1
            continue
        current.append(char)
        index += 1

    cells.append("".join(current).strip())
    if len(cells) < 2:
        return None
    return cells


def _parse_table_alignments(cells: list[str]) -> list[str] | None:
    aligns: list[str] = []
    for cell in cells:
        normalized = cell.replace(" ", "")
        if not re.fullmatch(r":?-{3,}:?", normalized):
            return None
        if normalized.startswith(":") and normalized.endswith(":"):
            aligns.append("center")
        elif normalized.endswith(":"):
            aligns.append("right")
        else:
            aligns.append("left")
    return aligns


def _normalize_table_cells(cells: list[str], column_count: int) -> list[str]:
    normalized = cells[:column_count]
    if len(normalized) < column_count:
        normalized.extend([""] * (column_count - len(normalized)))
    return normalized


def _parse_markdown_table(lines: list[str], start_index: int) -> tuple[TableBlock, int] | None:
    if start_index + 1 >= len(lines):
        return None

    header_cells = _split_table_cells(lines[start_index])
    separator_cells = _split_table_cells(lines[start_index + 1])
    if header_cells is None or separator_cells is None:
        return None
    if len(header_cells) != len(separator_cells):
        return None

    aligns = _parse_table_alignments(separator_cells)
    if aligns is None:
        return None

    column_count = len(header_cells)
    rows: list[list[str]] = []
    index = start_index + 2

    while index < len(lines):
        current = lines[index]
        if not current.strip():
            break
        row_cells = _split_table_cells(current)
        if row_cells is None:
            break
        rows.append(_normalize_table_cells(row_cells, column_count))
        index += 1

    return (
        TableBlock(
            headers=_normalize_table_cells(header_cells, column_count),
            aligns=aligns,
            rows=rows,
        ),
        index,
    )


def _parse_message_blocks(text: str) -> list[TextBlock | CodeBlock | TableBlock | MathBlock | ThematicBreak]:
    if not text:
        return []

    lines = text.splitlines()
    blocks: list[TextBlock | CodeBlock | TableBlock | MathBlock | ThematicBreak] = []
    paragraph: list[MarkdownLine] = []
    index = 0

    def flush_paragraph() -> None:
        if paragraph:
            blocks.append(TextBlock(lines=paragraph.copy()))
            paragraph.clear()

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if stripped.startswith("```"):
            flush_paragraph()
            code_lines: list[str] = []
            index += 1
            closed = False
            while index < len(lines):
                current = lines[index]
                if current.strip().startswith("```"):
                    closed = True
                    index += 1
                    break
                code_lines.append(current)
                index += 1
            if closed:
                blocks.append(CodeBlock(text="\n".join(code_lines)))
                continue
            paragraph.append(MarkdownLine(kind="paragraph", text=line))
            for current in code_lines:
                paragraph.append(MarkdownLine(kind="paragraph", text=current))
            break

        table_block = _parse_markdown_table(lines, index)
        if table_block is not None:
            flush_paragraph()
            table, next_index = table_block
            blocks.append(table)
            index = next_index
            continue

        single_line_block = _parse_single_line_block_math(line)
        if single_line_block is not None:
            flush_paragraph()
            blocks.append(single_line_block)
            index += 1
            continue

        if stripped == "$$":
            flush_paragraph()
            raw_lines = [line]
            latex_lines: list[str] = []
            index += 1
            closed = False
            while index < len(lines):
                current = lines[index]
                raw_lines.append(current)
                if current.strip() == "$$":
                    closed = True
                    index += 1
                    break
                latex_lines.append(current)
                index += 1
            if closed:
                blocks.append(MathBlock(latex="\n".join(latex_lines).strip(), raw_source="\n".join(raw_lines)))
                continue
            for current in raw_lines:
                paragraph.append(MarkdownLine(kind="paragraph", text=current))
            break

        if stripped == "":
            flush_paragraph()
            index += 1
            continue

        classified = _classify_markdown_line(line)
        if isinstance(classified, ThematicBreak):
            flush_paragraph()
            blocks.append(classified)
            index += 1
            continue

        if classified.kind in {"heading", "bullet", "ordered", "blockquote"}:
            flush_paragraph()
            blocks.append(TextBlock(lines=[classified]))
            index += 1
            continue

        paragraph.append(classified)
        index += 1

    flush_paragraph()
    return blocks


def _parse_streaming_message_blocks(text: str) -> ParsedStreamingBlocks:
    if not text:
        return ParsedStreamingBlocks([], 0)

    lines = text.splitlines()
    blocks: list[TextBlock | CodeBlock | TableBlock | MathBlock | ThematicBreak] = []
    paragraph: list[MarkdownLine] = []
    index = 0

    def flush_paragraph() -> None:
        if paragraph:
            blocks.append(TextBlock(lines=paragraph.copy()))
            paragraph.clear()

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if stripped.startswith("```"):
            flush_paragraph()
            code_lines: list[str] = []
            index += 1
            while index < len(lines):
                current = lines[index]
                if current.strip().startswith("```"):
                    index += 1
                    break
                code_lines.append(current)
                index += 1
            blocks.append(CodeBlock(text="\n".join(code_lines)))
            continue

        table_block = _parse_markdown_table(lines, index)
        if table_block is not None:
            flush_paragraph()
            table, next_index = table_block
            blocks.append(table)
            index = next_index
            continue

        single_line_block = _parse_single_line_block_math(line)
        if single_line_block is not None:
            flush_paragraph()
            blocks.append(single_line_block)
            index += 1
            continue

        if stripped == "$$":
            flush_paragraph()
            raw_lines = [line]
            latex_lines: list[str] = []
            index += 1
            closed = False
            while index < len(lines):
                current = lines[index]
                raw_lines.append(current)
                if current.strip() == "$$":
                    closed = True
                    index += 1
                    break
                latex_lines.append(current)
                index += 1
            if closed:
                blocks.append(MathBlock(latex="\n".join(latex_lines).strip(), raw_source="\n".join(raw_lines)))
            else:
                blocks.append(TextBlock(lines=[MarkdownLine(kind="paragraph", text=current) for current in raw_lines]))
            continue

        if stripped == "":
            flush_paragraph()
            index += 1
            continue

        classified = _classify_markdown_line(line)
        if isinstance(classified, ThematicBreak):
            flush_paragraph()
            blocks.append(classified)
            index += 1
            continue

        if classified.kind in {"heading", "bullet", "ordered", "blockquote"}:
            flush_paragraph()
            blocks.append(TextBlock(lines=[classified]))
            index += 1
            continue

        paragraph.append(classified)
        index += 1

    flush_paragraph()
    stable_count = max(0, len(blocks) - 1)
    if text.endswith("\n\n"):
        stable_count = len(blocks)
    return ParsedStreamingBlocks(blocks, stable_count)


def _stable_int_hash(text: str) -> int:
    return int(hashlib.sha1(text.encode("utf-8", errors="ignore")).hexdigest()[:12], 16)


def _block_fingerprint(block: TextBlock | CodeBlock | TableBlock | MathBlock | ThematicBreak) -> str:
    if isinstance(block, TextBlock):
        parts = [f"{line.kind}:{line.level}:{line.indent}:{line.marker}:{line.text}" for line in block.lines]
        return "text|" + "\n".join(parts)
    if isinstance(block, CodeBlock):
        return "code|" + block.text
    if isinstance(block, TableBlock):
        rows = ["\t".join(block.headers), "\t".join(block.aligns)]
        rows.extend("\t".join(row) for row in block.rows)
        return "table|" + "\n".join(rows)
    if isinstance(block, MathBlock):
        return "math|" + block.raw_source
    return "break|"


def _parse_inline_runs(text: str) -> list[TextRun | CodeRun | MathRun]:
    text = _normalize_inline_markdown_delimiters(text)
    runs: list[TextRun | CodeRun | MathRun] = []
    buffer: list[str] = []
    index = 0
    bold = False

    def flush_text() -> None:
        if buffer:
            runs.append(TextRun("".join(buffer), bold=bold))
            buffer.clear()

    while index < len(text):
        delimiter = _match_bold_delimiter(text, index)
        if delimiter is not None:
            flush_text()
            bold = not bold
            index += len(delimiter)
            continue

        char = text[index]

        if char == "`":
            end = text.find("`", index + 1)
            if end != -1:
                flush_text()
                runs.append(CodeRun(text[index + 1 : end]))
                index = end + 1
                continue

        if char == "$" and not _is_escaped(text, index):
            if index + 1 < len(text) and text[index + 1] == "$":
                buffer.append(char)
                index += 1
                continue
            end = _find_inline_math_end(text, index + 1)
            if end != -1:
                flush_text()
                raw_source = text[index : end + 1]
                runs.append(MathRun(text[index + 1 : end], raw_source))
                index = end + 1
                continue

        buffer.append(char)
        index += 1

    flush_text()
    return runs


def _find_inline_math_end(text: str, start: int) -> int:
    index = start
    while index < len(text):
        if text[index] == "$" and not _is_escaped(text, index):
            if index + 1 < len(text) and text[index + 1] == "$":
                return -1
            return index
        index += 1
    return -1


def _is_escaped(text: str, index: int) -> bool:
    backslashes = 0
    probe = index - 1
    while probe >= 0 and text[probe] == "\\":
        backslashes += 1
        probe -= 1
    return backslashes % 2 == 1


def _is_cjk(char: str) -> bool:
    code = ord(char)
    return (
        0x4E00 <= code <= 0x9FFF
        or 0x3400 <= code <= 0x4DBF
        or 0x3040 <= code <= 0x30FF
        or 0xAC00 <= code <= 0xD7AF
    )


def _tokenize_plain_text(text: str) -> list[str]:
    if not text:
        return []
    tokens: list[str] = []
    buffer: list[str] = []

    def flush_buffer() -> None:
        if buffer:
            tokens.append("".join(buffer))
            buffer.clear()

    for char in text:
        if char.isspace():
            if buffer:
                buffer.append(char)
            elif tokens:
                tokens[-1] += char
            else:
                buffer.append(char)
            continue
        if _is_cjk(char):
            flush_buffer()
            tokens.append(char)
            continue
        if char.isalnum() or char in {"_", "-", "/", "."}:
            buffer.append(char)
            continue
        flush_buffer()
        tokens.append(char)
    flush_buffer()
    return tokens


class Worker(QObject):
    loading_started = Signal()
    loading_finished = Signal()
    model_name_detected = Signal(str)
    error_raised = Signal(str)
    prefill_estimated = Signal(float, int)
    prefill_progress = Signal(int, int, float)
    prefill_finished = Signal()
    compression_started = Signal(str, str)
    chunk_received = Signal(str)
    turn_finished = Signal(object)
    context_cleared = Signal()

    def __init__(self) -> None:
        super().__init__()
        self.session = None

    def _ensure_session(self):
        if self.session is None:
            from chat_backend import ChatSession

            self.session = ChatSession()
        return self.session

    @Slot()
    def load_model(self) -> None:
        self.loading_started.emit()
        try:
            session = self._ensure_session()
            self.model_name_detected.emit(getattr(session, "model_display_name", HEADER_TITLE) or HEADER_TITLE)
            session.load()
        except Exception as exc:
            self.error_raised.emit(str(exc))
            return
        self.loading_finished.emit()

    @Slot(str, str)
    def submit_turn(self, text: str, image_path: str) -> None:
        prefill_pending = False
        try:
            from chat_backend import CompressionProgress, PrefillProgress, TurnResult, TurnStats, UserTurn

            session = self._ensure_session()
            turn = UserTurn(text=text, image_path=image_path or None)
            if session.should_show_backtracking_status(turn):
                self.compression_started.emit(GUI_TEXT.retrieving_context, "active")
            estimate = session.estimate_prefill(turn)
            self.prefill_estimated.emit(estimate.seconds, estimate.prompt_tokens)
            prefill_pending = True
            generator = session.stream_turn(turn)
            while True:
                try:
                    chunk = next(generator)
                except StopIteration as stop:
                    if prefill_pending:
                        self.prefill_finished.emit()
                        prefill_pending = False
                    result = stop.value or TurnResult(text="", stats=TurnStats())
                    self.turn_finished.emit(result)
                    break
                else:
                    if isinstance(chunk, PrefillProgress):
                        self.prefill_progress.emit(
                            chunk.prompt_tokens,
                            chunk.total_prompt_tokens,
                            chunk.prompt_tps,
                        )
                        continue
                    if isinstance(chunk, CompressionProgress):
                        if prefill_pending:
                            self.prefill_finished.emit()
                            prefill_pending = False
                        self.compression_started.emit(chunk.message, chunk.tone)
                        continue
                    if prefill_pending:
                        self.prefill_finished.emit()
                        prefill_pending = False
                    self.chunk_received.emit(chunk.text)
        except Exception as exc:
            if prefill_pending:
                self.prefill_finished.emit()
            self.error_raised.emit(str(exc))

    @Slot()
    def clear_context(self) -> None:
        session = self._ensure_session()
        session.clear()
        self.context_cleared.emit()


class Composer(QPlainTextEdit):
    submit_requested = Signal()

    def __init__(self) -> None:
        super().__init__()
        self.setObjectName("composer")
        self.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Minimum)
        self.setLineWrapMode(QPlainTextEdit.WidgetWidth)
        self.setWordWrapMode(QTextOption.WrapAtWordBoundaryOrAnywhere)
        self.setVerticalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        self.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        self._height_padding = 10
        self.document().contentsChanged.connect(self._sync_height)
        self._sync_height()

    def keyPressEvent(self, event) -> None:  # type: ignore[override]
        if event.key() in (Qt.Key_Return, Qt.Key_Enter):
            if event.modifiers() & Qt.ShiftModifier:
                super().keyPressEvent(event)
                return
            self.submit_requested.emit()
            event.accept()
            return
        super().keyPressEvent(event)

    def resizeEvent(self, event) -> None:  # type: ignore[override]
        super().resizeEvent(event)
        self._sync_height()

    def _measure_content_height(self) -> int:
        probe = QTextDocument(self)
        probe.setDefaultFont(self.font())
        probe.setDocumentMargin(self.document().documentMargin())

        option = probe.defaultTextOption()
        option.setWrapMode(QTextOption.WrapAtWordBoundaryOrAnywhere)
        probe.setDefaultTextOption(option)

        probe.setPlainText(self.toPlainText() or " ")
        probe.setTextWidth(max(1, self.viewport().width()))
        return int(probe.size().height())

    def _sync_height(self) -> None:
        line_height = self.fontMetrics().lineSpacing()
        min_height = max(38, line_height + self._height_padding)
        max_height = max(min_height, line_height * COMPOSER_MAX_VISIBLE_LINES + self._height_padding)
        content_height = self._measure_content_height() + self._height_padding
        target_height = max(min_height, min(max_height, content_height))
        self.setFixedHeight(target_height)
        if content_height > max_height:
            self.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        else:
            self.setVerticalScrollBarPolicy(Qt.ScrollBarAlwaysOff)


class LoadingSpinner(QWidget):
    def __init__(self, parent: QWidget | None = None) -> None:
        super().__init__(parent)
        self._step = 0
        self._spoke_count = 12
        self._timer = QTimer(self)
        self._timer.setInterval(90)
        self._timer.timeout.connect(self._advance)
        self.setFixedSize(41, 41)

    def start(self) -> None:
        if not self._timer.isActive():
            self._timer.start()
        self.show()
        self.update()

    def stop(self) -> None:
        self._timer.stop()
        self.update()

    def _advance(self) -> None:
        self._step = (self._step + 1) % self._spoke_count
        self.update()

    def paintEvent(self, event) -> None:  # type: ignore[override]
        del event
        painter = QPainter(self)
        painter.setRenderHint(QPainter.Antialiasing)
        painter.translate(self.rect().center())

        outer_radius = 14
        inner_radius = 8
        pen_width = 3

        for index in range(self._spoke_count):
            distance = (index - self._step) % self._spoke_count
            opacity = 1.0 - (distance / self._spoke_count) * 0.72
            color = QColor(160, 160, 160)
            color.setAlphaF(max(0.22, opacity))
            painter.save()
            painter.rotate(index * (360 / self._spoke_count))
            painter.setPen(QPen(color, pen_width, Qt.SolidLine, Qt.RoundCap))
            painter.drawLine(0, -inner_radius, 0, -outer_radius)
            painter.restore()


class PrefillProgressLine(QWidget):
    def __init__(self) -> None:
        super().__init__()
        self._progress = 0.0
        self._estimate_seconds = 0.0
        self._start_time = 0.0
        self._complete_from = 0.0
        self._complete_started_at = 0.0
        self._active = False
        self._completing = False
        self._hold_until = 0.0
        self._timer = QTimer(self)
        self._timer.setInterval(16)
        self._timer.timeout.connect(self._tick)
        self.setObjectName("prefillProgressLine")
        self.setFixedHeight(2)
        self.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)

    def start_estimate(self, seconds: float) -> None:
        self._estimate_seconds = max(0.45, float(seconds))
        self._start_time = time.monotonic()
        self._active = True
        self._completing = False
        self._hold_until = 0.0
        self._complete_from = self._progress if self._progress > 0 else 0.0
        if not self._timer.isActive():
            self._timer.start()
        self.update()

    def sync_runtime(self, progress_ratio: float, total_seconds: float) -> None:
        total_seconds = max(0.45, float(total_seconds))
        progress_ratio = max(0.0, min(1.0, float(progress_ratio)))
        now = time.monotonic()
        self._estimate_seconds = total_seconds
        self._start_time = now - min(progress_ratio, 0.95) * total_seconds
        self._active = True
        self._completing = False
        self._hold_until = 0.0
        self._progress = max(self._progress, min(0.95, 0.95 * progress_ratio))
        if not self._timer.isActive():
            self._timer.start()
        self.update()

    def complete(self) -> None:
        if not self._active and self._progress <= 0 and not self._completing:
            return
        self._active = False
        self._completing = True
        self._complete_from = max(self._progress, 0.12)
        self._complete_started_at = time.monotonic()
        self._hold_until = 0.0
        if not self._timer.isActive():
            self._timer.start()
        self.update()

    def finish_now(self) -> None:
        self._active = False
        self._completing = False
        self._progress = 1.0
        self._hold_until = time.monotonic() + 0.08
        if not self._timer.isActive():
            self._timer.start()
        self.update()

    def reset(self) -> None:
        self._active = False
        self._completing = False
        self._hold_until = 0.0
        self._progress = 0.0
        self._timer.stop()
        self.update()

    def _tick(self) -> None:
        now = time.monotonic()
        if self._active:
            elapsed = max(0.0, now - self._start_time)
            estimate = max(0.01, self._estimate_seconds)
            if elapsed <= estimate:
                target = 0.95 * min(1.0, elapsed / estimate)
            else:
                overflow = elapsed - estimate
                overflow_window = max(0.5, estimate * 0.6)
                overflow_ratio = min(1.0, overflow / overflow_window)
                target = 0.95 + 0.035 * overflow_ratio
            self._progress = max(self._progress, min(0.985, target))
            self.update()
            return

        if self._completing:
            duration = 0.18
            ratio = min(1.0, (now - self._complete_started_at) / duration)
            eased = 1.0 - (1.0 - ratio) ** 3
            self._progress = self._complete_from + (1.0 - self._complete_from) * eased
            self.update()
            if ratio >= 1.0:
                self._completing = False
                self._hold_until = now + 0.08
            return

        if self._hold_until > 0.0:
            if now >= self._hold_until:
                self._hold_until = 0.0
                self._progress = 0.0
                self._timer.stop()
                self.update()
            return

        self._timer.stop()

    def paintEvent(self, event) -> None:  # type: ignore[override]
        del event
        painter = QPainter(self)
        painter.fillRect(self.rect(), QColor(0, 0, 0, 20))
        if self._progress <= 0:
            return
        fill_width = max(1, int(round(self.width() * min(1.0, self._progress))))
        painter.fillRect(0, 0, fill_width, self.height(), QColor(34, 197, 94))


class MessageBubble(QFrame):
    def __init__(
        self,
        role: str,
        text: str = "",
        image_path: str | None = None,
        math_renderer: MathRenderService | None = None,
    ) -> None:
        super().__init__()
        self.role = role
        self._raw_text = text
        self._streaming = False
        self._math_renderer = math_renderer

        self.setObjectName("userBubble" if role == "user" else "assistantBubble")
        self.setSizePolicy(QSizePolicy.Maximum, QSizePolicy.Maximum)
        if role == "user":
            self.setMaximumWidth(680)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(14, 12, 14, 12)
        layout.setSpacing(8)

        self.image_label: QLabel | None = None
        if image_path:
            self.image_label = QLabel()
            self.image_label.setAlignment(Qt.AlignCenter)
            self.image_label.setScaledContents(False)
            self.image_label.setMaximumHeight(260)
            pixmap = QPixmap(image_path)
            if not pixmap.isNull():
                scaled = pixmap.scaled(
                    300,
                    220,
                    Qt.KeepAspectRatio,
                    Qt.SmoothTransformation,
                )
                self.image_label.setPixmap(scaled)
                layout.addWidget(self.image_label)

        self.body_wrap = QWidget()
        if role == "user":
            self.body_wrap.setSizePolicy(QSizePolicy.Maximum, QSizePolicy.Fixed)
        else:
            self.body_wrap.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
        self.body_layout = QVBoxLayout(self.body_wrap)
        self.body_layout.setContentsMargins(0, 0, 0, 0)
        self.body_layout.setSpacing(0)

        self.user_label: QLabel | None = None
        self.assistant_view: AssistantContentWidget | None = None
        self._assistant_render_timer: QTimer | None = None
        self._assistant_layout_sync_pending = False

        if role == "user":
            self.user_label = QLabel()
            self.user_label.setObjectName("userBody")
            self.user_label.setTextFormat(Qt.PlainText)
            self.user_label.setWordWrap(True)
            self.user_label.setAlignment(Qt.AlignLeft | Qt.AlignVCenter)
            self.user_label.setMargin(0)
            self.user_label.setContentsMargins(0, 0, 0, 0)
            self.user_label.setSizePolicy(QSizePolicy.Maximum, QSizePolicy.Fixed)
            self.body_layout.addWidget(self.user_label)
        else:
            self.assistant_view = AssistantContentWidget(math_renderer or MathRenderService())
            self.assistant_view.content_height_changed.connect(self._schedule_assistant_body_sync)
            self.body_layout.addWidget(self.assistant_view)
            self._assistant_render_timer = QTimer(self)
            self._assistant_render_timer.setSingleShot(True)
            self._assistant_render_timer.setInterval(ASSISTANT_RENDER_INTERVAL_MS)
            self._assistant_render_timer.timeout.connect(self._render_assistant_content)

        layout.addWidget(self.body_wrap)
        self.set_text(text)

    def refresh_layout_for_width(self) -> None:
        if self.role == "assistant":
            self._sync_assistant_body_size()
            self._schedule_assistant_body_sync()
        else:
            self._sync_user_body_size()
            QTimer.singleShot(0, self._sync_user_body_size)

    def _render_assistant_content(self) -> None:
        if self.role != "assistant" or self.assistant_view is None:
            return
        self.assistant_view.render_text(self._raw_text or "")
        self._sync_assistant_body_size()
        self._schedule_assistant_body_sync()

    def _schedule_assistant_render(self, force: bool = False) -> None:
        if self.role != "assistant" or self.assistant_view is None:
            return
        if force or self._assistant_render_timer is None:
            self._render_assistant_content()
            return
        if not self._assistant_render_timer.isActive():
            self._assistant_render_timer.start()

    def _notify_layout_change(self) -> None:
        self.body_wrap.updateGeometry()
        self.updateGeometry()
        parent = self.parentWidget()
        while parent is not None:
            layout = parent.layout()
            if layout is not None:
                layout.invalidate()
            parent.updateGeometry()
            parent = parent.parentWidget()

    def _schedule_assistant_body_sync(self) -> None:
        if self._assistant_layout_sync_pending:
            return
        self._assistant_layout_sync_pending = True
        QTimer.singleShot(0, self._run_assistant_body_sync)

    def _run_assistant_body_sync(self) -> None:
        self._assistant_layout_sync_pending = False
        self._sync_assistant_body_size()

    def _sync_assistant_body_size(self) -> None:
        if self.assistant_view is None:
            return
        margins = self.layout().contentsMargins()
        available_width = max(1, self.width() - margins.left() - margins.right())
        self.body_wrap.setFixedWidth(available_width)
        self.assistant_view.setFixedWidth(available_width)
        self.assistant_view.updateGeometry()
        height = self.assistant_view.heightForWidth(available_width)
        if height <= 0:
            height = self.assistant_view.sizeHint().height()
        height = max(1, height)
        if self.body_wrap.width() == available_width and self.assistant_view.height() == height and self.body_wrap.height() == height:
            return
        self.assistant_view.setFixedHeight(height)
        self.body_wrap.setFixedHeight(height)
        self._notify_layout_change()

    def _sync_user_body_size(self) -> None:
        if self.role != "user" or self.user_label is None:
            return

        self.ensurePolished()
        self.user_label.ensurePolished()

        margins = self.layout().contentsMargins()
        max_text_width = max(1, self.maximumWidth() - margins.left() - margins.right())
        text = self._raw_text or " "

        self.user_label.setText(text)
        self.user_label.setWordWrap(False)
        single_line_width = max(18, self.user_label.sizeHint().width())
        self.user_label.setWordWrap(True)

        text_width = min(max_text_width, single_line_width)
        self.user_label.setFixedWidth(text_width)
        self.user_label.updateGeometry()

        text_height = self.user_label.heightForWidth(text_width)
        if text_height <= 0:
            text_height = self.user_label.sizeHint().height()
        text_height = max(self.user_label.fontMetrics().height(), text_height) + 4

        self.user_label.setFixedHeight(text_height)
        self.body_wrap.setFixedWidth(text_width)
        self.body_wrap.setFixedHeight(text_height)
        self._notify_layout_change()

    def begin_stream(self) -> None:
        if self._streaming:
            return
        self._streaming = True
        if self.role == "assistant" and self.assistant_view is not None:
            self.assistant_view.begin_stream()
        if self._assistant_render_timer is not None:
            self._assistant_render_timer.setInterval(STREAMING_RENDER_INTERVAL_MS)

    def end_stream(self) -> None:
        was_streaming = self._streaming
        self._streaming = False
        if self._assistant_render_timer is not None:
            self._assistant_render_timer.setInterval(ASSISTANT_RENDER_INTERVAL_MS)
        if self.role == "assistant" and was_streaming:
            if self.assistant_view is not None:
                self.assistant_view.finish_stream(self._raw_text or "")
                self._schedule_assistant_body_sync()

    def append_text(self, chunk: str, immediate: bool = False) -> None:
        self.begin_stream()
        if self.role == "assistant":
            self._raw_text += chunk
            if self.assistant_view is not None:
                self.assistant_view.apply_stream_units([chunk])
                self._schedule_assistant_body_sync()
            return
        self.set_text(self._raw_text + chunk)

    def apply_stream_units(self, units: list[str]) -> None:
        self.begin_stream()
        if self.role != "assistant":
            self.set_text(self._raw_text + "".join(units))
            return
        self._raw_text += "".join(units)
        if self.assistant_view is not None:
            self.assistant_view.apply_stream_units(units)
            self._schedule_assistant_body_sync()

    def set_pending(self) -> None:
        self.begin_stream()
        self.set_text(PENDING_TEXT)

    def set_text(self, text: str) -> None:
        self._raw_text = text
        if self.role == "assistant":
            if self.assistant_view is None:
                return
            if self._streaming:
                self._schedule_assistant_render(force=text == PENDING_TEXT)
            else:
                self._schedule_assistant_render(force=True)
        else:
            if self.user_label is None:
                return
            self.user_label.setText(text or "")
            self._sync_user_body_size()
            QTimer.singleShot(0, self._sync_user_body_size)


class CopyMessageButton(QPushButton):
    def __init__(self) -> None:
        super().__init__()
        self._copied = False
        self.setObjectName("copyMessageButton")
        self.setToolTip(GUI_TEXT.copy_message_tooltip)
        self.setAccessibleName(GUI_TEXT.copy_message_tooltip)
        self.setCursor(Qt.PointingHandCursor)
        self.setFocusPolicy(Qt.NoFocus)
        self.setFixedSize(28, 28)
        self.setText("")

    def set_copied(self, copied: bool) -> None:
        self._copied = copied
        self.update()

    def paintEvent(self, event) -> None:  # type: ignore[override]
        del event
        painter = QPainter(self)
        painter.setRenderHint(QPainter.Antialiasing, True)

        if self.isDown():
            painter.setBrush(QColor(0, 0, 0, 22))
        elif self.underMouse() and self.isEnabled():
            painter.setBrush(QColor(0, 0, 0, 12))
        else:
            painter.setBrush(Qt.NoBrush)
        painter.setPen(Qt.NoPen)
        painter.drawRoundedRect(QRectF(2, 2, 24, 24), 6, 6)

        if self._copied:
            painter.setBrush(Qt.NoBrush)
            painter.setPen(QPen(QColor(34, 197, 94), 2.0, Qt.SolidLine, Qt.RoundCap, Qt.RoundJoin))
            painter.drawLine(8, 15, 12, 19)
            painter.drawLine(12, 19, 21, 9)
            return

        icon_color = QColor(93, 93, 93) if self.isEnabled() else QColor(180, 180, 180)
        painter.setBrush(Qt.NoBrush)
        painter.setPen(QPen(icon_color, 1.45, Qt.SolidLine, Qt.RoundCap, Qt.RoundJoin))
        painter.drawRoundedRect(QRectF(10, 5, 13, 13), 3, 3)
        painter.drawRoundedRect(QRectF(6, 10, 13, 13), 3, 3)


class MessageRow(QWidget):
    def __init__(
        self,
        role: str,
        text: str = "",
        image_path: str | None = None,
        math_renderer: MathRenderService | None = None,
    ) -> None:
        super().__init__()
        self.role = role
        self.bubble = MessageBubble(role=role, text=text, image_path=image_path, math_renderer=math_renderer)
        self.copy_button = CopyMessageButton()
        self.copy_button.clicked.connect(self.copy_to_clipboard)
        self._copy_tooltip_reset_timer = QTimer(self)
        self._copy_tooltip_reset_timer.setSingleShot(True)
        self._copy_tooltip_reset_timer.timeout.connect(self._reset_copy_tooltip)

        self.setObjectName("assistantRow" if role == "assistant" else "userRow")
        self.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Maximum)

        layout = QHBoxLayout(self)
        layout.setContentsMargins(20, 0, 20, 0)
        layout.setSpacing(0)

        self.message_stack = QWidget()
        self.message_stack.setObjectName("messageStack")
        stack_layout = QVBoxLayout(self.message_stack)
        stack_layout.setContentsMargins(0, 0, 0, 0)
        stack_layout.setSpacing(2)
        stack_layout.addWidget(self.bubble, 0, Qt.AlignTop)
        stack_layout.addWidget(
            self.copy_button,
            0,
            (Qt.AlignLeft if role == "assistant" else Qt.AlignRight) | Qt.AlignTop,
        )

        if role == "assistant":
            layout.addStretch(1)
            layout.addWidget(self.message_stack, 0, Qt.AlignTop)
            layout.addStretch(1)
        else:
            layout.addStretch(1)
            layout.addWidget(self.message_stack, 0, Qt.AlignRight | Qt.AlignTop)

        self._update_copy_button_state()

    def resizeEvent(self, event) -> None:  # type: ignore[override]
        super().resizeEvent(event)
        if self.role == "assistant":
            self.bubble.setFixedWidth(max(320, int(self.width() * 0.75)))
            self.bubble.refresh_layout_for_width()
            self.message_stack.setFixedWidth(self.bubble.width())

    @property
    def _raw_text(self) -> str:
        return self.bubble._raw_text

    def _update_copy_button_state(self) -> None:
        text = self.bubble._raw_text or ""
        has_copyable_text = bool(text.strip()) and text != PENDING_TEXT
        visible = has_copyable_text and not (self.role == "assistant" and self.bubble._streaming)
        self.copy_button.setVisible(visible)
        self.copy_button.setEnabled(visible)

    def _reset_copy_tooltip(self) -> None:
        self.copy_button.setToolTip(GUI_TEXT.copy_message_tooltip)
        self.copy_button.set_copied(False)

    @Slot()
    def copy_to_clipboard(self) -> None:
        text = self.bubble._raw_text or ""
        if not text.strip():
            return
        QApplication.clipboard().setText(text)
        self.copy_button.setToolTip(GUI_TEXT.copied_message_tooltip)
        self.copy_button.set_copied(True)
        self._copy_tooltip_reset_timer.start(10000)

    def append_text(self, chunk: str, immediate: bool = False) -> None:
        self.bubble.append_text(chunk, immediate=immediate)
        self._update_copy_button_state()

    def apply_stream_units(self, units: list[str]) -> None:
        self.bubble.apply_stream_units(units)
        self._update_copy_button_state()

    def set_pending(self) -> None:
        self.bubble.set_pending()
        self._update_copy_button_state()

    def set_text(self, text: str) -> None:
        self.bubble.set_text(text)
        self._update_copy_button_state()

    def end_stream(self) -> None:
        self.bubble.end_stream()
        self._update_copy_button_state()


class MainWindow(QMainWindow):
    load_requested = Signal()
    submit_requested = Signal(str, str)
    clear_requested = Signal()

    def __init__(self) -> None:
        super().__init__()
        self.setWindowTitle(WINDOW_TITLE)
        self.resize(980, 760)
        self.math_renderer = MathRenderService(self)

        self._loaded = False
        self._busy = True
        self._current_reply: MessageRow | None = None
        self._attached_image: str | None = None
        self._attached_file: FileAttachment | None = None
        self._stick_to_bottom = True
        self._auto_scroll_threshold = AUTO_SCROLL_BOTTOM_THRESHOLD
        self._suppress_scroll_tracking = False
        self._manual_scroll_pending = False
        self._welcome_added = False
        self._prefill_prompt_tokens = 0
        self._prefill_locked = False
        self._auto_scroll_timer = QTimer(self)
        self._auto_scroll_timer.setSingleShot(True)
        self._auto_scroll_timer.setTimerType(Qt.PreciseTimer)
        self._auto_scroll_timer.timeout.connect(self._scroll_to_bottom)
        self._display_flush_timer = QTimer(self)
        self._display_flush_timer.setInterval(DISPLAY_STREAM_INTERVAL_MS)
        self._display_flush_timer.setTimerType(Qt.PreciseTimer)
        self._display_flush_timer.timeout.connect(self._flush_display_queue)
        self._display_queue: list[str] = []
        self._pending_turn_result = None
        self._streaming_auto_scroll_pending = False
        self._scroll_animation_timer = QTimer(self)
        self._scroll_animation_timer.setInterval(SCROLL_FRAME_INTERVAL_MS)
        self._scroll_animation_timer.setTimerType(Qt.PreciseTimer)
        self._scroll_animation_timer.timeout.connect(self._advance_scroll_animation)
        self._scroll_target_value: float | None = None
        self._scroll_animation_mode: str | None = None
        self._status_delay_timer = QTimer(self)
        self._status_delay_timer.setSingleShot(True)
        self._status_delay_timer.setTimerType(Qt.PreciseTimer)
        self._status_delay_timer.timeout.connect(self._apply_deferred_status)
        self._active_status_hold_until = 0.0
        self._deferred_status: tuple[str, str] | None = None

        self._build_ui()
        self._apply_styles()
        self._set_loading_ui(True)
        self._update_attachment_label()
        self._update_controls()
        self._set_status(GUI_TEXT.loading_model, force=True)

    def shutdown(self) -> None:
        self.math_renderer.shutdown()

    def bind_worker(self, worker: Worker, thread: QThread) -> None:
        self.load_requested.connect(worker.load_model)
        self.submit_requested.connect(worker.submit_turn)
        self.clear_requested.connect(worker.clear_context)

        worker.loading_started.connect(self.on_loading_started)
        worker.loading_finished.connect(self.on_loading_finished)
        worker.model_name_detected.connect(self.on_model_name_detected)
        worker.error_raised.connect(self.on_error)
        worker.prefill_estimated.connect(self.on_prefill_estimated)
        worker.prefill_progress.connect(self.on_prefill_progress)
        worker.prefill_finished.connect(self.on_prefill_finished)
        worker.compression_started.connect(self.on_compression_started)
        worker.chunk_received.connect(self.on_chunk_received)
        worker.turn_finished.connect(self.on_turn_finished)
        worker.context_cleared.connect(self.on_context_cleared)

        thread.finished.connect(worker.deleteLater)

    @Slot(str)
    def on_model_name_detected(self, model_name: str) -> None:
        self.header_title.setText(model_name.strip() or HEADER_TITLE)

    def _build_ui(self) -> None:
        root = QWidget()
        root.setObjectName("root")
        outer = QVBoxLayout(root)
        outer.setContentsMargins(0, 0, 0, 0)
        outer.setSpacing(0)

        self.header = QWidget()
        self.header.setObjectName("header")
        header_layout = QHBoxLayout(self.header)
        header_layout.setContentsMargins(14, 12, 14, 12)
        header_layout.setSpacing(0)

        header_left_spacer = QWidget()
        header_left_spacer.setFixedWidth(36)
        header_layout.addWidget(header_left_spacer)

        header_center = QWidget()
        header_center_layout = QVBoxLayout(header_center)
        header_center_layout.setContentsMargins(0, 0, 0, 0)
        header_center_layout.setSpacing(2)

        self.header_title = QLabel(HEADER_TITLE)
        self.header_title.setObjectName("headerTitle")
        self.header_title.setAlignment(Qt.AlignCenter)
        header_center_layout.addWidget(self.header_title)

        self.header_status = QLabel("")
        self.header_status.setObjectName("headerStatus")
        self.header_status.setAlignment(Qt.AlignCenter)
        header_center_layout.addWidget(self.header_status)

        header_layout.addWidget(header_center, 1)

        self.new_chat_button = QPushButton()
        self.new_chat_button.setObjectName("newChatButton")
        self.new_chat_button.setToolTip(GUI_TEXT.new_chat_tooltip)
        self.new_chat_button.setIcon(QIcon(str(NEW_CHAT_ICON)))
        self.new_chat_button.setIconSize(QSize(20, 20))
        self.new_chat_button.setFixedSize(36, 36)
        self.new_chat_button.clicked.connect(self.request_clear)
        header_layout.addWidget(self.new_chat_button, 0, Qt.AlignRight | Qt.AlignVCenter)

        outer.addWidget(self.header)

        self.scroll_area = QScrollArea()
        self.scroll_area.setFrameShape(QFrame.NoFrame)
        self.scroll_area.setWidgetResizable(True)
        self.scroll_area.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        self.scroll_area.installEventFilter(self)
        scroll_bar = self.scroll_area.verticalScrollBar()
        scroll_bar.valueChanged.connect(self._on_scroll_value_changed)
        scroll_bar.rangeChanged.connect(self._on_scroll_range_changed)
        scroll_bar.sliderPressed.connect(self._on_scroll_slider_pressed)
        scroll_bar.sliderReleased.connect(self._on_scroll_slider_released)
        scroll_bar.actionTriggered.connect(self._on_scroll_action_triggered)
        scroll_bar.installEventFilter(self)
        self.scroll_area.viewport().installEventFilter(self)

        self.chat_container = QWidget()
        self.chat_container.setObjectName("chatContainer")
        self.chat_layout = QVBoxLayout(self.chat_container)
        self.chat_layout.setContentsMargins(0, 20, 0, 20)
        self.chat_layout.setSpacing(10)
        self.chat_layout.setAlignment(Qt.AlignTop)

        self.scroll_area.setWidget(self.chat_container)
        outer.addWidget(self.scroll_area, 1)

        self.input_wrap = QWidget()
        self.input_wrap.setObjectName("inputWrap")
        input_wrap_layout = QVBoxLayout(self.input_wrap)
        input_wrap_layout.setContentsMargins(20, 18, 20, 20)
        input_wrap_layout.setSpacing(0)

        center_row = QHBoxLayout()
        center_row.setContentsMargins(0, 0, 0, 0)
        center_row.setSpacing(0)
        center_row.addStretch(1)

        self.input_panel = QFrame()
        self.input_panel.setObjectName("inputPanel")

        shadow = QGraphicsDropShadowEffect(self.input_panel)
        shadow.setBlurRadius(24)
        shadow.setOffset(0, 4)
        shadow.setColor(QColor(0, 0, 0, 24))
        self.input_panel.setGraphicsEffect(shadow)

        input_panel_layout = QVBoxLayout(self.input_panel)
        input_panel_layout.setContentsMargins(12, 9, 12, 9)
        input_panel_layout.setSpacing(6)

        self.attachment_label = QLabel()
        self.attachment_label.setObjectName("attachmentLabel")
        input_panel_layout.addWidget(self.attachment_label)

        input_row = QHBoxLayout()
        input_row.setContentsMargins(0, 0, 0, 0)
        input_row.setSpacing(8)

        self.pick_attachment_button = QPushButton("+")
        self.pick_attachment_button.setObjectName("plusButton")
        self.pick_attachment_button.clicked.connect(self.pick_attachment)
        self.pick_attachment_button.setFixedSize(32, 32)
        input_row.addWidget(self.pick_attachment_button)

        self.remove_attachment_button = QPushButton(GUI_TEXT.remove_image)
        self.remove_attachment_button.setObjectName("toolButton")
        self.remove_attachment_button.clicked.connect(self.remove_attachment)
        input_row.addWidget(self.remove_attachment_button)

        self.composer = Composer()
        self.composer.setPlaceholderText(PLACEHOLDER)
        self.composer.submit_requested.connect(self.request_submit)
        self.composer.textChanged.connect(self._update_controls)
        input_row.addWidget(self.composer, 1)

        self.send_button = QPushButton("➤")
        self.send_button.setObjectName("sendButton")
        self.send_button.clicked.connect(self.request_submit)
        self.send_button.setFixedSize(32, 32)
        input_row.addWidget(self.send_button)

        input_panel_layout.addLayout(input_row)
        center_row.addWidget(self.input_panel)
        center_row.addStretch(1)

        input_wrap_layout.addLayout(center_row)

        self.loading_wrap = QWidget()
        self.loading_wrap.setObjectName("loadingWrap")
        loading_wrap_layout = QVBoxLayout(self.loading_wrap)
        loading_wrap_layout.setContentsMargins(20, 24, 20, 28)
        loading_wrap_layout.setSpacing(0)
        loading_wrap_layout.addStretch(1)

        spinner_row = QHBoxLayout()
        spinner_row.setContentsMargins(0, 0, 0, 0)
        spinner_row.setSpacing(0)
        spinner_row.addStretch(1)
        self.loading_spinner = LoadingSpinner()
        spinner_row.addWidget(self.loading_spinner, 0, Qt.AlignHCenter | Qt.AlignVCenter)
        spinner_row.addStretch(1)
        loading_wrap_layout.addLayout(spinner_row)
        loading_wrap_layout.addStretch(1)

        self.bottom_stack_host = QWidget()
        self.bottom_stack_host.setObjectName("bottomStackHost")
        bottom_host_layout = QVBoxLayout(self.bottom_stack_host)
        bottom_host_layout.setContentsMargins(0, 0, 0, 0)
        bottom_host_layout.setSpacing(0)

        self.prefill_progress_line = PrefillProgressLine()
        bottom_host_layout.addWidget(self.prefill_progress_line)

        self.bottom_stack_widget = QWidget()
        self.bottom_stack = QStackedLayout(self.bottom_stack_widget)
        self.bottom_stack.setContentsMargins(0, 0, 0, 0)
        self.bottom_stack.setStackingMode(QStackedLayout.StackOne)
        self.bottom_stack.addWidget(self.loading_wrap)
        self.bottom_stack.addWidget(self.input_wrap)
        bottom_host_layout.addWidget(self.bottom_stack_widget)
        outer.addWidget(self.bottom_stack_host)

        self.setCentralWidget(root)
        self._sync_input_panel_width()

    def _apply_styles(self) -> None:
        self.setStyleSheet(
            """
            QMainWindow, QWidget#root {
                background: #FFFFFF;
                color: #37352F;
                font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", "Helvetica Neue", Arial, sans-serif;
            }
            QWidget#header {
                background: #FFFFFF;
                border-bottom: 1px solid rgba(0, 0, 0, 0.05);
            }
            QLabel#headerTitle {
                color: #6B6B6B;
                font-size: 15px;
                font-weight: 500;
            }
            QLabel#headerStatus {
                color: #9B9B9B;
                font-size: 12px;
            }
            QLabel#headerStatus[statusTone="compression"] {
                color: #D97706;
                font-weight: 600;
            }
            QLabel#headerStatus[statusTone="active"] {
                color: #16A34A;
                font-weight: 600;
            }
            QScrollArea {
                border: none;
                background: #FFFFFF;
            }
            QWidget#chatContainer {
                background: #FFFFFF;
            }
            QWidget#bottomStackHost {
                border-top: none;
            }
            QPushButton#newChatButton {
                background: transparent;
                border: none;
                border-radius: 8px;
            }
            QPushButton#newChatButton:hover {
                background: #F3F3F3;
            }
            QPushButton#newChatButton:disabled {
                background: transparent;
            }
            QWidget#assistantRow {
                background: transparent;
            }
            QWidget#userRow {
                background: transparent;
            }
            QFrame#assistantBubble {
                background: transparent;
                border: none;
            }
            QWidget#assistantParagraphBlock {
                background: transparent;
            }
            QWidget#assistantLine_heading,
            QWidget#assistantLine_paragraph,
            QWidget#assistantLine_bullet,
            QWidget#assistantLine_ordered,
            QWidget#assistantLine_blockquote {
                background: transparent;
            }
            QFrame#userBubble {
                background: #1677FF;
                border: none;
                border-radius: 12px;
            }
            QLabel#assistantTextToken {
                background: transparent;
                color: #37352F;
                font-size: 15px;
                line-height: 1.45;
            }
            QLabel#assistantTextToken[markdownBold="true"] {
                font-weight: 700;
            }
            QWidget#assistantHeadingFlow[headingLevel="1"] QLabel#assistantTextToken,
            QWidget#assistantHeadingFlow[headingLevel="1"] QLabel#assistantMathFallbackInline {
                color: #222222;
                font-size: 24px;
                font-weight: 700;
            }
            QWidget#assistantHeadingFlow[headingLevel="2"] QLabel#assistantTextToken,
            QWidget#assistantHeadingFlow[headingLevel="2"] QLabel#assistantMathFallbackInline {
                color: #222222;
                font-size: 21px;
                font-weight: 700;
            }
            QWidget#assistantHeadingFlow[headingLevel="3"] QLabel#assistantTextToken,
            QWidget#assistantHeadingFlow[headingLevel="3"] QLabel#assistantMathFallbackInline {
                color: #222222;
                font-size: 18px;
                font-weight: 700;
            }
            QWidget#assistantHeadingFlow[headingLevel="4"] QLabel#assistantTextToken,
            QWidget#assistantHeadingFlow[headingLevel="5"] QLabel#assistantTextToken,
            QWidget#assistantHeadingFlow[headingLevel="6"] QLabel#assistantTextToken,
            QWidget#assistantHeadingFlow[headingLevel="4"] QLabel#assistantMathFallbackInline,
            QWidget#assistantHeadingFlow[headingLevel="5"] QLabel#assistantMathFallbackInline,
            QWidget#assistantHeadingFlow[headingLevel="6"] QLabel#assistantMathFallbackInline {
                color: #222222;
                font-size: 16px;
                font-weight: 700;
            }
            QLabel#assistantListMarker,
            QLabel#assistantQuoteMarker {
                background: transparent;
                color: #6B6B6B;
                font-size: 15px;
                font-weight: 600;
            }
            QLabel#assistantCodeToken {
                background: #F7F7F8;
                color: #37352F;
                font-size: 14px;
                font-family: ui-monospace, SFMono-Regular, Menlo, Monaco, Consolas, monospace;
                border-radius: 6px;
                padding: 1px 5px;
            }
            QLabel#assistantMathFallbackInline {
                background: transparent;
                color: #7A4E00;
                font-size: 14px;
            }
            QLabel#assistantMathFallbackBlock {
                background: #FFF7E6;
                color: #7A4E00;
                font-size: 14px;
                border-radius: 8px;
                padding: 8px 10px;
            }
            QFrame#assistantCodeBlock {
                background: #F7F7F8;
                border: none;
                border-radius: 8px;
            }
            QFrame#assistantTable {
                background: rgba(0, 0, 0, 0.08);
                border: 1px solid rgba(0, 0, 0, 0.08);
                border-radius: 10px;
            }
            QWidget#assistantTableRow {
                background: transparent;
            }
            QFrame#assistantTableCell {
                background: #FFFFFF;
                border: none;
            }
            QFrame#assistantTableCell[tableHeader="true"] {
                background: #F7F7F8;
            }
            QFrame#assistantTableCell[alternateRow="true"] {
                background: #FBFBFC;
            }
            QLabel#assistantCodeBlockLabel {
                background: transparent;
                color: #37352F;
                font-size: 14px;
                font-family: ui-monospace, SFMono-Regular, Menlo, Monaco, Consolas, monospace;
            }
            QLabel#userBody {
                background: transparent;
                color: #FFFFFF;
                font-size: 15px;
            }
            QWidget#inputWrap {
                background: qlineargradient(
                    x1: 0, y1: 0, x2: 0, y2: 1,
                    stop: 0 rgba(255, 255, 255, 0),
                    stop: 0.45 rgba(255, 255, 255, 255),
                    stop: 1 rgba(255, 255, 255, 255)
                );
            }
            QWidget#loadingWrap {
                background: qlineargradient(
                    x1: 0, y1: 0, x2: 0, y2: 1,
                    stop: 0 rgba(255, 255, 255, 0),
                    stop: 0.45 rgba(255, 255, 255, 255),
                    stop: 1 rgba(255, 255, 255, 255)
                );
            }
            QFrame#inputPanel {
                background: #FFFFFF;
                border: 1px solid rgba(0, 0, 0, 0.05);
                border-radius: 14px;
            }
            QLabel#attachmentLabel {
                color: #6B6B6B;
                font-size: 12px;
                padding-left: 2px;
            }
            QPushButton#plusButton {
                min-width: 32px;
                min-height: 32px;
                max-width: 32px;
                max-height: 32px;
                color: #6B6B6B;
                background: transparent;
                border: 1px solid rgba(0, 0, 0, 0.08);
                border-radius: 8px;
                font-size: 16px;
                font-weight: 500;
            }
            QPushButton#plusButton:hover {
                background: #F7F7F8;
            }
            QPushButton#toolButton {
                background: transparent;
                color: #6B6B6B;
                border: none;
                padding: 0 4px;
                font-size: 13px;
            }
            QPushButton#toolButton:hover {
                color: #37352F;
            }
            QPlainTextEdit#composer {
                background: transparent;
                border: none;
                color: #2B2F36;
                font-size: 16px;
                padding: 5px 0;
            }
            QPlainTextEdit#composer[readOnly="true"] {
                color: #A0A0A0;
            }
            QPushButton#sendButton {
                background: #1677FF;
                color: #FFFFFF;
                border: none;
                border-radius: 8px;
                font-size: 14px;
                font-weight: 600;
            }
            QPushButton#sendButton:hover {
                background: #0F63D6;
            }
            QPushButton#sendButton:disabled {
                background: #CCCCCC;
                color: #FFFFFF;
            }
            QScrollBar:vertical {
                width: 10px;
                background: transparent;
                margin: 8px 0 8px 0;
            }
            QScrollBar::handle:vertical {
                background: rgba(0, 0, 0, 0.16);
                border-radius: 5px;
                min-height: 36px;
            }
            QScrollBar::add-line:vertical,
            QScrollBar::sub-line:vertical,
            QScrollBar::add-page:vertical,
            QScrollBar::sub-page:vertical {
                background: none;
                border: none;
                height: 0px;
            }
            """
        )

    def _apply_status_immediately(self, text: str, tone: str = "normal") -> None:
        self.header_status.setProperty("statusTone", tone)
        self.header_status.style().unpolish(self.header_status)
        self.header_status.style().polish(self.header_status)
        self.header_status.setText(text)

    def _set_status(self, text: str, tone: str = "normal", *, force: bool = False) -> None:
        now = time.monotonic()
        if force:
            self._deferred_status = None
            self._active_status_hold_until = 0.0
            self._status_delay_timer.stop()
            self._apply_status_immediately(text, tone)
            return

        if tone == "active":
            self._deferred_status = None
            self._active_status_hold_until = now + ACTIVE_STATUS_MIN_VISIBLE_MS / 1000.0
            self._status_delay_timer.stop()
            self._apply_status_immediately(text, tone)
            return

        if now < self._active_status_hold_until:
            self._deferred_status = (text, tone)
            delay_ms = max(1, math.ceil((self._active_status_hold_until - now) * 1000))
            self._status_delay_timer.start(delay_ms)
            return

        self._deferred_status = None
        self._active_status_hold_until = 0.0
        self._apply_status_immediately(text, tone)

    @Slot()
    def _apply_deferred_status(self) -> None:
        if self._deferred_status is None:
            return
        now = time.monotonic()
        if now < self._active_status_hold_until:
            delay_ms = max(1, math.ceil((self._active_status_hold_until - now) * 1000))
            self._status_delay_timer.start(delay_ms)
            return

        text, tone = self._deferred_status
        self._deferred_status = None
        self._active_status_hold_until = 0.0
        self._apply_status_immediately(text, tone)

    def _reset_prefill_progress(self) -> None:
        self._prefill_prompt_tokens = 0
        self._prefill_locked = False
        self.prefill_progress_line.reset()

    def _reset_display_queue(self) -> None:
        self._display_flush_timer.stop()
        self._display_queue.clear()
        self._pending_turn_result = None

    def _sync_input_panel_width(self) -> None:
        host_width = self.centralWidget().width() if self.centralWidget() is not None else self.width()
        available_width = max(480, host_width - 40)
        target_width = max(520, int(available_width * INPUT_PANEL_WIDTH_RATIO))
        self.input_panel.setFixedWidth(min(available_width, target_width))

    def _set_loading_ui(self, loading: bool) -> None:
        if loading:
            self.bottom_stack.setCurrentWidget(self.loading_wrap)
            self.loading_spinner.start()
        else:
            self.loading_spinner.stop()
            self.bottom_stack.setCurrentWidget(self.input_wrap)

    def _add_message(self, role: str, text: str = "", image_path: str | None = None) -> MessageRow:
        widget = MessageRow(role=role, text=text, image_path=image_path, math_renderer=self.math_renderer)
        self.chat_layout.addWidget(widget)
        self._request_auto_scroll()
        return widget

    def _add_welcome_message(self) -> None:
        self._welcome_added = True
        self._add_message("assistant", text=WELCOME_TEXT)

    def _clear_messages(self) -> None:
        while self.chat_layout.count() > 0:
            item = self.chat_layout.takeAt(0)
            widget = item.widget()
            if widget is not None:
                widget.deleteLater()

    def _request_auto_scroll(self) -> None:
        if not self._stick_to_bottom:
            return
        if self._busy and self._current_reply is not None:
            self._streaming_auto_scroll_pending = True
            if not self._auto_scroll_timer.isActive():
                self._auto_scroll_timer.start(DISPLAY_STREAM_INTERVAL_MS)
            return
        if not self._auto_scroll_timer.isActive():
            self._auto_scroll_timer.start(SCROLL_FRAME_INTERVAL_MS)

    def _scroll_to_bottom(self) -> None:
        bar = self.scroll_area.verticalScrollBar()
        if self._streaming_auto_scroll_pending:
            self._streaming_auto_scroll_pending = False
            self._scroll_target_value = None
            self._scroll_animation_mode = None
            if self._scroll_animation_timer.isActive():
                self._scroll_animation_timer.stop()
            self._set_scrollbar_value(bar, bar.maximum(), suppress_tracking=True)
            return
        self._animate_scroll_to(bar.maximum(), mode="auto")

    def _begin_manual_scroll(self, moving_up: bool = False) -> None:
        self._manual_scroll_pending = True
        self._auto_scroll_timer.stop()
        self._streaming_auto_scroll_pending = False
        if moving_up:
            self._stick_to_bottom = False

    def eventFilter(self, watched: QObject, event: QEvent) -> bool:  # type: ignore[override]
        if hasattr(self, "scroll_area"):
            scroll_bar = self.scroll_area.verticalScrollBar()
            viewport = self.scroll_area.viewport()
            if watched is scroll_bar or watched is viewport or watched is self.scroll_area:
                if event.type() == QEvent.Wheel:
                    pixel_delta_getter = getattr(event, "pixelDelta", None)
                    angle_delta_getter = getattr(event, "angleDelta", None)
                    pixel_delta = pixel_delta_getter().y() if callable(pixel_delta_getter) else 0
                    angle_delta = angle_delta_getter().y() if callable(angle_delta_getter) else 0
                    if pixel_delta == 0 and angle_delta == 0:
                        return super().eventFilter(watched, event)
                    self._begin_manual_scroll(moving_up=(pixel_delta > 0 or angle_delta > 0))
                    self._enqueue_manual_scroll_delta(pixel_delta, angle_delta)
                    event.accept()
                    return True
        return super().eventFilter(watched, event)

    def _enqueue_manual_scroll_delta(self, pixel_delta: int, angle_delta: int) -> None:
        bar = self.scroll_area.verticalScrollBar()
        if self._scroll_target_value is None:
            self._scroll_target_value = float(bar.value())

        if pixel_delta != 0:
            delta = float(pixel_delta)
        else:
            line_height = max(18, self.fontMetrics().lineSpacing())
            delta = (angle_delta / 120.0) * line_height * 3.0

        target = self._scroll_target_value - delta
        self._scroll_target_value = max(float(bar.minimum()), min(float(bar.maximum()), target))
        self._scroll_animation_mode = "manual"
        if not self._scroll_animation_timer.isActive():
            self._scroll_animation_timer.start()

    def _animate_scroll_to(self, value: int, mode: str) -> None:
        bar = self.scroll_area.verticalScrollBar()
        self._scroll_target_value = max(float(bar.minimum()), min(float(bar.maximum()), float(value)))
        self._scroll_animation_mode = mode
        if not self._scroll_animation_timer.isActive():
            self._scroll_animation_timer.start()

    @Slot()
    def _advance_scroll_animation(self) -> None:
        if self._scroll_target_value is None:
            self._scroll_animation_timer.stop()
            self._scroll_animation_mode = None
            return

        bar = self.scroll_area.verticalScrollBar()
        current = float(bar.value())
        distance = self._scroll_target_value - current

        if abs(distance) <= SCROLL_MIN_STEP:
            next_value = int(round(self._scroll_target_value))
            mode = self._scroll_animation_mode
            self._scroll_target_value = None
            self._scroll_animation_mode = None
            self._set_scrollbar_value(bar, next_value, suppress_tracking=(mode == "auto"))
            self._scroll_animation_timer.stop()
            return

        step = distance * SCROLL_ANIMATION_FACTOR
        if abs(step) < SCROLL_MIN_STEP:
            step = SCROLL_MIN_STEP if distance > 0 else -SCROLL_MIN_STEP

        next_value = current + step
        if distance > 0:
            next_value = min(next_value, self._scroll_target_value)
        else:
            next_value = max(next_value, self._scroll_target_value)

        self._set_scrollbar_value(
            bar,
            int(round(next_value)),
            suppress_tracking=(self._scroll_animation_mode == "auto"),
        )

    def _set_scrollbar_value(self, bar, value: int, *, suppress_tracking: bool) -> None:
        if suppress_tracking:
            self._suppress_scroll_tracking = True
        try:
            bar.setValue(value)
        finally:
            if suppress_tracking:
                self._suppress_scroll_tracking = False

    def _drain_display_queue_now(self) -> None:
        if self._current_reply is not None and self._display_queue:
            units = self._display_queue.copy()
            self._display_queue.clear()
            self._current_reply.apply_stream_units(units)
        self._display_flush_timer.stop()

    def _consume_display_units(self) -> int:
        if not self._display_queue:
            return 0
        backlog = len(self._display_queue)
        budget = min(
            DISPLAY_STREAM_MAX_UNITS,
            max(DISPLAY_STREAM_MIN_UNITS, 1 + backlog // DISPLAY_STREAM_BACKLOG_DIVISOR),
        )
        if self._current_reply is None:
            self._display_queue.clear()
            return 0
        units = self._display_queue[:budget]
        del self._display_queue[:budget]
        self._current_reply.apply_stream_units(units)
        return sum(len(unit) for unit in units)

    @Slot()
    def _flush_display_queue(self) -> None:
        consumed = self._consume_display_units()
        if consumed == 0:
            self._display_flush_timer.stop()
        if not self._display_queue and self._pending_turn_result is not None:
            result = self._pending_turn_result
            self._pending_turn_result = None
            self._finalize_turn_result(result)

    def _is_near_bottom(self) -> bool:
        bar = self.scroll_area.verticalScrollBar()
        return bar.maximum() - bar.value() <= self._auto_scroll_threshold

    def _on_scroll_value_changed(self, value: int) -> None:
        del value
        if self._suppress_scroll_tracking:
            return
        near_bottom = self._is_near_bottom()
        if self._manual_scroll_pending or self.scroll_area.verticalScrollBar().isSliderDown():
            self._stick_to_bottom = near_bottom
            self._manual_scroll_pending = False
            if not near_bottom:
                self._auto_scroll_timer.stop()
            return
        if not near_bottom and self._stick_to_bottom:
            self._stick_to_bottom = False
            self._auto_scroll_timer.stop()

    def _on_scroll_range_changed(self, minimum: int, maximum: int) -> None:
        del minimum, maximum
        self._request_auto_scroll()

    @Slot()
    def _on_scroll_slider_pressed(self) -> None:
        self._begin_manual_scroll()

    @Slot()
    def _on_scroll_slider_released(self) -> None:
        self._manual_scroll_pending = True
        self._on_scroll_value_changed(self.scroll_area.verticalScrollBar().value())

    @Slot(int)
    def _on_scroll_action_triggered(self, action: int) -> None:
        moving_up = action in {
            QAbstractSlider.SliderSingleStepSub,
            QAbstractSlider.SliderPageStepSub,
            QAbstractSlider.SliderToMinimum,
        }
        self._begin_manual_scroll(moving_up=moving_up)

    def _set_busy(self, value: bool) -> None:
        self._busy = value
        self._update_controls()

    def _update_controls(self) -> None:
        ready = self._loaded and not self._busy
        can_edit = self._loaded
        has_image = bool(self._attached_image)
        has_file = self._attached_file is not None
        has_attachment = has_image or has_file
        has_content = bool(self.composer.toPlainText().strip() or has_attachment)
        self.send_button.setEnabled(ready and has_content)
        self.new_chat_button.setEnabled(self._loaded and not self._busy)
        self.pick_attachment_button.setEnabled(not self._busy)
        self.remove_attachment_button.setEnabled(has_attachment and not self._busy)
        self.composer.setReadOnly(not can_edit)

    def _update_attachment_label(self) -> None:
        has_image = bool(self._attached_image)
        has_file = self._attached_file is not None
        has_attachment = has_image or has_file
        self.attachment_label.setVisible(has_attachment)
        self.remove_attachment_button.setVisible(has_attachment)
        if has_image:
            self.attachment_label.setText(GUI_TEXT.selected_image.format(name=Path(self._attached_image).name))
        elif self._attached_file is not None:
            self.attachment_label.setText(GUI_TEXT.selected_file.format(name=self._attached_file.name))
        else:
            self.attachment_label.clear()
        self._update_controls()

    def pick_attachment(self) -> None:
        document_filter = "Documents (*.md *.txt *.docx *.doc *.xlsx *.xls *.csv *.html)"
        image_filter = "Images (*.png *.jpg *.jpeg *.webp *.bmp *.gif)"
        path, _ = QFileDialog.getOpenFileName(
            self,
            GUI_TEXT.choose_image_title,
            "",
            f"{document_filter};;{image_filter}",
        )
        if not path:
            return

        extension = Path(path).suffix.lower()
        if extension in SUPPORTED_IMAGE_EXTENSIONS:
            self._attached_image = path
            self._attached_file = None
            self._update_attachment_label()
            return
        if extension not in SUPPORTED_DOCUMENT_EXTENSIONS:
            QMessageBox.critical(self, GUI_TEXT.error_title, GUI_TEXT.unsupported_file)
            return
        try:
            attachment = read_file_attachment(path)
        except Exception as exc:
            QMessageBox.critical(self, GUI_TEXT.error_title, GUI_TEXT.file_read_failed.format(error=exc))
            return
        self._attached_image = None
        self._attached_file = attachment
        self._update_attachment_label()

    def pick_image(self) -> None:
        self.pick_attachment()

    def remove_attachment(self) -> None:
        self._attached_image = None
        self._attached_file = None
        self._update_attachment_label()

    def remove_image(self) -> None:
        self.remove_attachment()

    def request_submit(self) -> None:
        if not self._loaded or self._busy:
            return

        text = self.composer.toPlainText().strip()
        image_path = self._attached_image
        file_attachment = self._attached_file
        if not text and not image_path and file_attachment is None:
            return

        backend_text = text
        display_text = text
        if file_attachment is not None:
            file_label = GUI_TEXT.file_fallback.format(name=file_attachment.name)
            display_text = f"{text}\n\n{file_label}".strip() if text else file_label
            file_content = file_attachment.text
            if file_attachment.truncated:
                file_content = f"{file_content}{GUI_TEXT.file_prompt_truncated}"
            file_block = f"{GUI_TEXT.file_prompt_header.format(name=file_attachment.name)}\n{file_content}"
            backend_text = f"{text}\n\n{file_block}".strip() if text else f"{GUI_TEXT.file_prompt_intro}\n\n{file_block}"

        self._stick_to_bottom = True
        self._add_message("user", text=display_text or GUI_TEXT.image_fallback, image_path=image_path)
        self._current_reply = self._add_message("assistant", text="")
        self._current_reply.set_pending()

        self.composer.clear()
        self._attached_image = None
        self._attached_file = None
        self._update_attachment_label()
        self._reset_display_queue()
        self._reset_prefill_progress()
        self._set_busy(True)
        self._set_status(GUI_TEXT.generating_answer, force=True)
        self.submit_requested.emit(backend_text, image_path or "")

    def request_clear(self) -> None:
        if not self._loaded or self._busy:
            return
        self.clear_requested.emit()

    @Slot()
    def on_loading_started(self) -> None:
        self._set_busy(True)
        self._reset_display_queue()
        self._reset_prefill_progress()
        self._set_loading_ui(True)
        self._set_status(GUI_TEXT.loading_model, force=True)

    @Slot()
    def on_loading_finished(self) -> None:
        self._loaded = True
        self._set_busy(False)
        self._reset_display_queue()
        self._reset_prefill_progress()
        self._set_loading_ui(False)
        self._set_status(GUI_TEXT.model_loaded, force=True)

    @Slot(float, int)
    def on_prefill_estimated(self, seconds: float, prompt_tokens: int) -> None:
        if self._prefill_locked:
            return
        self._prefill_prompt_tokens = max(0, int(prompt_tokens))
        self.prefill_progress_line.start_estimate(seconds)
        self._set_status(GUI_TEXT.prefill_estimate.format(seconds=seconds))

    @Slot(int, int, float)
    def on_prefill_progress(self, prompt_tokens: int, total_prompt_tokens: int, prompt_tps: float) -> None:
        if self._prefill_locked:
            return
        if total_prompt_tokens <= 0 or prompt_tps <= 0:
            return
        progress_ratio = max(0.0, min(1.0, prompt_tokens / total_prompt_tokens))
        total_seconds = total_prompt_tokens / prompt_tps
        self._prefill_prompt_tokens = total_prompt_tokens
        self.prefill_progress_line.sync_runtime(progress_ratio, total_seconds)
        self._set_status(GUI_TEXT.prefill_estimate.format(seconds=total_seconds))

    @Slot()
    def on_prefill_finished(self) -> None:
        if self._prefill_locked:
            return
        self.prefill_progress_line.complete()
        if self._busy:
            self._set_status(GUI_TEXT.generating_answer)

    @Slot(str)
    def on_chunk_received(self, chunk: str) -> None:
        if not self._prefill_locked:
            self._prefill_locked = True
            self.prefill_progress_line.finish_now()
        if self._current_reply is None:
            self._current_reply = self._add_message("assistant", text="")
        if self._current_reply._raw_text == PENDING_TEXT:
            self._current_reply.set_text("")
        self._display_queue.extend(_split_display_units(chunk))
        if not self._display_flush_timer.isActive():
            self._display_flush_timer.start()

    @Slot(str, str)
    def on_compression_started(self, message: str, tone: str) -> None:
        if (
            self._current_reply is not None
            and self._current_reply._raw_text
            and self._current_reply._raw_text != PENDING_TEXT
        ):
            self._drain_display_queue_now()
            self._current_reply.end_stream()
        self._set_status(message or GUI_TEXT.compressing_context, tone=tone or "compression")

    def _finalize_turn_result(self, result) -> None:
        if self._current_reply is None:
            self._current_reply = self._add_message("assistant", text=result.text)
            self._current_reply.end_stream()
        else:
            self._current_reply.set_text(result.text)
            self._current_reply.end_stream()
        self._current_reply = None
        self._set_busy(False)
        self._request_auto_scroll()
        self._set_status(
            GUI_TEXT.completed_status.format(
                prompt=result.stats.prompt_tokens,
                generation=result.stats.generation_tokens,
                peak=result.stats.peak_memory,
            )
        )

    @Slot(object)
    def on_turn_finished(self, result) -> None:
        if self._display_queue:
            self._pending_turn_result = result
            if not self._display_flush_timer.isActive():
                self._display_flush_timer.start()
            return
        self._finalize_turn_result(result)
        self._prefill_prompt_tokens = 0

    @Slot()
    def on_context_cleared(self) -> None:
        self._reset_display_queue()
        self._clear_messages()
        self._current_reply = None
        self._stick_to_bottom = True
        self._welcome_added = False
        self._add_welcome_message()
        self._reset_prefill_progress()
        self._set_status(GUI_TEXT.context_cleared, force=True)

    @Slot(str)
    def on_error(self, message: str) -> None:
        self._reset_display_queue()
        if self._current_reply is not None and self._current_reply._raw_text == PENDING_TEXT:
            self._current_reply.set_text(GUI_TEXT.generation_failed)
        self._current_reply = None
        self._set_busy(False)
        self._reset_prefill_progress()
        if not self._loaded:
            self.loading_spinner.stop()
        self._set_status(GUI_TEXT.error_status, force=True)
        QMessageBox.critical(self, GUI_TEXT.error_title, message)

    def closeEvent(self, event: QCloseEvent) -> None:  # type: ignore[override]
        self.shutdown()
        event.accept()

    def resizeEvent(self, event) -> None:  # type: ignore[override]
        super().resizeEvent(event)
        self._sync_input_panel_width()

    def showEvent(self, event) -> None:  # type: ignore[override]
        super().showEvent(event)
        if not self._welcome_added:
            QTimer.singleShot(0, self._add_welcome_message)


def run_app() -> int:
    app = QApplication(sys.argv)
    window = MainWindow()

    thread = QThread()
    worker = Worker()
    worker.moveToThread(thread)
    window.bind_worker(worker, thread)

    thread.start()
    window.show()
    window.load_requested.emit()

    exit_code = app.exec()
    window.shutdown()
    thread.quit()
    thread.wait()
    return exit_code
